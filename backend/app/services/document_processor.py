# backend/app/services/document_processor.py
# ⚠️ REMPLACER SEULEMENT LA MÉTHODE apply_global_replacements ET AJOUTER L'IMPORT

import io
import logging
import re  # ✅ AJOUTER CETTE LIGNE
from typing import BinaryIO, Tuple, Dict, List
from docx import Document
from docx.shared import Inches
import PyPDF2
import pytesseract
from pdf2image import convert_from_bytes
from PIL import Image
import tempfile
import os
from .replacement_engine import IntelligentReplacementEngine

logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self):
        self.supported_extensions = ['.pdf', '.docx']

    def process_uploaded_file(self, file_content: bytes, filename: str) -> Tuple[Document, str]:
        """Traite un fichier uploadé et renvoie un Document et son texte."""
        extension = os.path.splitext(filename)[1].lower()

        if extension == '.docx':
            return self._process_docx(file_content)
        elif extension == '.pdf':
            return self._process_pdf(file_content)
        else:
            raise ValueError(f"Format de fichier non supporté: {extension}")

    def _process_pdf(self, pdf_content: bytes) -> Tuple[Document, str]:
        """Traite un PDF en tentant d'extraire le texte puis en OCR si besoin."""
        text = self._extract_text_from_pdf(pdf_content)
        if len(text.strip()) < 100:
            text = self._ocr_pdf_with_tesseract(pdf_content)
        doc = self._create_docx_from_text(text)
        return doc, text

    def _process_docx(self, docx_content: bytes) -> Tuple[Document, str]:
        """Charge un DOCX et en extrait le texte."""
        stream = io.BytesIO(docx_content)
        document = Document(stream)
        text = self._extract_text_from_docx(document)
        return document, text

    def _extract_text_from_pdf(self, pdf_content: bytes) -> str:
        text = ""
        try:
            reader = PyPDF2.PdfReader(io.BytesIO(pdf_content))
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        except Exception as e:
            logger.error(f"Erreur lors de l'extraction PDF: {e}")
        return text

    def _ocr_pdf_with_tesseract(self, pdf_content: bytes) -> str:
        text = ""
        try:
            images = convert_from_bytes(pdf_content, dpi=300)
            for image in images:
                page_text = pytesseract.image_to_string(image, lang='fra', config='--psm 6')
                text += page_text + "\n"
        except Exception as e:
            logger.error(f"Erreur lors de l'OCR: {e}")
        return text

    def _extract_text_from_docx(self, document: Document) -> str:
        parts = []
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        return "\n".join(parts)

    def _create_docx_from_text(self, text: str) -> Document:
        document = Document()
        document.add_heading('Document converti', 0)
        for line in text.split('\n'):
            if line.strip():
                document.add_paragraph(line.strip())
        return document

    # ... GARDER TOUTES LES AUTRES MÉTHODES INCHANGÉES ...
    
    def apply_global_replacements(self, document: Document, replacements: dict) -> bytes:
        """
        🔧 REPLACEMENT ENGINE CORRIGÉ - Résout "Saïd OULHADJ" → "X X"
        """
        try:
            logger.info(f"🔄 Application de {len(replacements)} remplacements")
            
            # Replacements dans les paragraphes
            for paragraph in document.paragraphs:
                if paragraph.text.strip():
                    self._replace_in_paragraph_fixed(paragraph, replacements)
            
            # Replacements dans les tableaux
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if paragraph.text.strip():
                                self._replace_in_paragraph_fixed(paragraph, replacements)
            
            # Sauvegarder en mémoire
            output_stream = io.BytesIO()
            document.save(output_stream)
            output_stream.seek(0)
            
            logger.info("✅ Remplacements appliqués avec succès")
            return output_stream.read()
            
        except Exception as e:
            logger.error(f"Erreur lors des remplacements: {e}")
            raise ValueError(f"Impossible d'appliquer les remplacements: {e}")

    def _replace_in_paragraph_fixed(self, paragraph, replacements: Dict[str, str]):
        """
        🎯 SOLUTION AU PROBLÈME : Remplacement complet sans fragmentation
        """
        try:
            full_text = paragraph.text
            if not full_text.strip():
                return
            
            # Appliquer tous les remplacements
            modified_text = full_text
            has_changes = False
            
            for original, replacement in replacements.items():
                if original.lower() in modified_text.lower():
                    # Remplacement insensible à la casse
                    pattern = re.compile(re.escape(original), re.IGNORECASE)
                    new_text = pattern.sub(replacement, modified_text)
                    if new_text != modified_text:
                        modified_text = new_text
                        has_changes = True
                        logger.debug(f"✅ '{original}' → '{replacement}'")
            
            if has_changes:
                # SOLUTION : Vider tous les runs et mettre le texte dans le premier
                for run in paragraph.runs:
                    run.text = ""
                
                if not paragraph.runs:
                    paragraph.add_run("")
                
                paragraph.runs[0].text = modified_text

        except Exception as e:
            logger.debug(f"Erreur replacement: {e}")

    def apply_intelligent_replacements(self, document: Document, entities_data: List[Dict],
                                       groups_data: List[Dict] = None) -> bytes:
        """Applique les remplacements en évitant les conflits."""
        try:
            logger.info("🚀 Début de l'anonymisation intelligente")

            engine = IntelligentReplacementEngine()

            if groups_data:
                for group in groups_data:
                    group_replacement = group.get('replacement', 'GROUPE_X')
                    entity_ids = group.get('entity_ids', [])

                    for entity_data in entities_data:
                        if entity_data.get('selected') and entity_data.get('id') in entity_ids:
                            engine.add_replacement(
                                original=entity_data['text'],
                                replacement=group_replacement,
                                entity_id=entity_data['id'],
                                is_grouped=True,
                                group_id=group.get('id')
                            )
                            logger.info(f"Groupe ajouté: '{entity_data['text']}' -> '{group_replacement}'")

            for entity_data in entities_data:
                if entity_data.get('selected') and not entity_data.get('is_grouped', False):
                    engine.add_replacement(
                        original=entity_data['text'],
                        replacement=entity_data['replacement'],
                        entity_id=entity_data['id'],
                        is_grouped=False
                    )
                    logger.info(f"Entité ajoutée: '{entity_data['text']}' -> '{entity_data['replacement']}'")

            self._apply_replacements_to_document(document, engine)

            report = engine.get_replacement_report()
            logger.info(
                f"📊 Rapport d'anonymisation: {report['total_rules']} règles, {report['active_rules']} actives, {report['grouped_rules']} groupées")

            output_stream = io.BytesIO()
            document.save(output_stream)
            output_stream.seek(0)

            logger.info("✅ Anonymisation intelligente terminée")
            return output_stream.read()

        except Exception as e:
            logger.error(f"Erreur lors de l'anonymisation intelligente: {e}")
            raise ValueError(f"Impossible d'appliquer l'anonymisation: {e}")

    def _apply_replacements_to_document(self, document: Document, engine: IntelligentReplacementEngine):
        """Applique les remplacements à toutes les parties du document."""
        replacements_count = 0

        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                original_text = paragraph.text
                new_text, stats = engine.apply_replacements(original_text)

                if new_text != original_text:
                    self._replace_paragraph_text(paragraph, new_text)
                    replacements_count += sum(stats.values())

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            original_text = paragraph.text
                            new_text, stats = engine.apply_replacements(original_text)

                            if new_text != original_text:
                                self._replace_paragraph_text(paragraph, new_text)
                                replacements_count += sum(stats.values())

        logger.info(f"✅ {replacements_count} remplacements appliqués dans le document")

    def _replace_paragraph_text(self, paragraph, new_text: str):
        """Remplace le texte d'un paragraphe en préservant le formatage."""
        for run in paragraph.runs:
            run.text = ""

        if paragraph.runs:
            paragraph.runs[0].text = new_text
        else:
            paragraph.add_run(new_text)

# Instance globale inchangée
document_processor = DocumentProcessor()
