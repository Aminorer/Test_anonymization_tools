import io
from docx import Document
from app.services.document_processor import DocumentProcessor


def test_create_docx_from_text():
    dp = DocumentProcessor()
    text = "Hello\nWorld"
    doc = dp._create_docx_from_text(text)
    paragraphs = [p.text for p in doc.paragraphs]
    assert paragraphs[0] == "Document converti"
    assert paragraphs[1] == "Hello"
    assert paragraphs[2] == "World"


def test_extract_text_from_docx():
    doc = Document()
    doc.add_paragraph("Para1")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Cell"
    dp = DocumentProcessor()
    text = dp._extract_text_from_docx(doc)
    assert "Para1" in text
    assert "Cell" in text


def test_apply_global_replacements():
    doc = Document()
    doc.add_paragraph("Hello NAME")
    dp = DocumentProcessor()
    result_bytes = dp.apply_global_replacements(doc, {"NAME": "John"})
    new_doc = Document(io.BytesIO(result_bytes))
    assert new_doc.paragraphs[0].text == "Hello John"


def test_process_uploaded_file_docx():
    original = Document()
    original.add_paragraph("Sample")
    stream = io.BytesIO()
    original.save(stream)
    stream.seek(0)
    dp = DocumentProcessor()
    doc, text = dp.process_uploaded_file(stream.read(), "test.docx")
    assert text.strip() == "Sample"

