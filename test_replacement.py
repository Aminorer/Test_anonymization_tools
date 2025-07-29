#!/usr/bin/env python3
"""
🧪 Test Simple du Replacement Engine v2.0
Vérifie que "Saïd OULHADJ" → "PERSONNE_A" fonctionne
"""

import io
from docx import Document

def test_replacement_fix():
    print("🧪 Test du Replacement Engine v2.0")
    print("=" * 50)
    
    # 1. Créer un document test
    doc = Document()
    doc.add_paragraph("Maître Saïd OULHADJ représente la partie.")
    doc.add_paragraph("Le dossier est traité par Saïd OULHADJ.")
    
    print("📄 Document original créé")
    for i, para in enumerate(doc.paragraphs):
        print(f"  Paragraphe {i+1}: {para.text}")
    
    # 2. Simuler le replacement avec la nouvelle méthode
    replacements = {
        "Saïd OULHADJ": "PERSONNE_A",
        "Maître": "TITRE_X"
    }
    
    print(f"\n🔄 Application des remplacements: {replacements}")
    
    # 3. Appliquer la méthode corrigée
    for paragraph in doc.paragraphs:
        _replace_in_paragraph_v2(paragraph, replacements)
    
    # 4. Vérifier le résultat
    print("\n✅ Document après remplacement:")
    for i, para in enumerate(doc.paragraphs):
        print(f"  Paragraphe {i+1}: {para.text}")
    
    # 5. Vérifications
    full_text = "\n".join([p.text for p in doc.paragraphs])
    
    # Tests
    success = True
    if "Saïd OULHADJ" in full_text:
        print("❌ ERREUR: 'Saïd OULHADJ' encore présent")
        success = False
    else:
        print("✅ 'Saïd OULHADJ' correctement remplacé")
    
    if "PERSONNE_A" in full_text:
        print("✅ 'PERSONNE_A' présent")
    else:
        print("❌ ERREUR: 'PERSONNE_A' manquant")
        success = False
    
    if "TITRE_X" in full_text:
        print("✅ 'TITRE_X' présent")
    else:
        print("❌ ERREUR: 'TITRE_X' manquant")
        success = False
    
    print(f"\n🎯 Résultat: {'SUCCÈS' if success else 'ÉCHEC'}")
    return success

def _replace_in_paragraph_v2(paragraph, replacements):
    """Version corrigée du replacement"""
    import re
    
    full_text = paragraph.text
    if not full_text.strip():
        return
    
    # Vérifier si on a des remplacements à faire
    has_replacements = False
    modified_text = full_text
    
    for original, replacement in replacements.items():
        if original.lower() in full_text.lower():
            has_replacements = True
            # Remplacement insensible à la casse
            pattern = re.compile(re.escape(original), re.IGNORECASE)
            modified_text = pattern.sub(replacement, modified_text)
            print(f"  🔄 '{original}' → '{replacement}'")
    
    if not has_replacements:
        return
    
    # SOLUTION ROBUSTE : Remplacer tout le contenu du paragraphe
    # Effacer tous les runs existants
    for run in paragraph.runs:
        run.text = ""
    
    # Si on n'a plus de runs, en créer un
    if not paragraph.runs:
        paragraph.add_run("")
    
    # Mettre le texte modifié dans le premier run
    paragraph.runs[0].text = modified_text

if __name__ == "__main__":
    test_replacement_fix()