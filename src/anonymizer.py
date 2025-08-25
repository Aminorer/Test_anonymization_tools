# src/anonymizer.py - VERSION NER FONCTIONNELLE CORRIGÉE
"""
Anonymiseur de Documents avec NER fonctionnel
Résolution des conflits PyTorch/Streamlit
"""

# === CONFIGURATION ANTI-CONFLIT ===
import os
import warnings
import logging
import threading
import uuid
from datetime import datetime
import json

# Configuration logging précoce
logging.basicConfig(level=logging.WARNING)
for logger_name in ["transformers", "torch", "tensorflow", "urllib3"]:
    logging.getLogger(logger_name).setLevel(logging.ERROR)

# Variables d'environnement critiques
os.environ.update({
    "TOKENIZERS_PARALLELISM": "false",
    "OMP_NUM_THREADS": "1",
    "MKL_NUM_THREADS": "1",
    "OPENBLAS_NUM_THREADS": "1",
    "VECLIB_MAXIMUM_THREADS": "1",
    "NUMEXPR_NUM_THREADS": "1",
    "KMP_DUPLICATE_LIB_OK": "TRUE",
    "PYTORCH_JIT": "0",
    "PYTORCH_JIT_USE_NNC": "0"
})

# Suppression warnings
warnings.filterwarnings("ignore", category=UserWarning)
warnings.filterwarnings("ignore", category=FutureWarning)
warnings.filterwarnings("ignore", category=DeprecationWarning)

# === IMPORTS STANDARDS ===
import re
import unicodedata
import tempfile
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple, Set
from typing import TYPE_CHECKING
from dataclasses import dataclass, asdict
from io import BytesIO
import hashlib
from .utils import (
    generate_anonymization_stats,
    serialize_entity_mapping,
    compute_confidence,
    get_name_normalization_titles,
    get_similarity_threshold,
    get_similarity_weights,
)
from .legal_normalizer import LegalEntityNormalizer
try:
    from rapidfuzz.distance import Levenshtein as RFLevenshtein
except ImportError:  # pragma: no cover - rapidfuzz is optional at runtime
    RFLevenshtein = None  # type: ignore
from .bktree import BKTree
if TYPE_CHECKING:  # pragma: no cover
    from .entity_manager import EntityManager

# === IMPORTS DOCUMENT ===
try:
    from docx import Document
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    logging.warning("Support DOCX désactivé: Module python-docx non disponible")

try:
    from pdf2docx import parse as pdf2docx_parse
    PDF2DOCX_SUPPORT = True
except ImportError:
    PDF2DOCX_SUPPORT = False
    logging.warning("Support PDF2DOCX désactivé: Module pdf2docx non disponible")

# === CONFIGURATION PYTORCH THREAD-SAFE ===
_pytorch_lock = threading.Lock()
_pytorch_configured = False

def configure_pytorch_safe():
    """Configuration PyTorch thread-safe pour Streamlit"""
    global _pytorch_configured
    
    if _pytorch_configured:
        return True
    
    with _pytorch_lock:
        if _pytorch_configured:
            return True
        
        try:
            import torch
            
            # Configuration threads de base
            try:
                torch.set_num_threads(1)
            except (RuntimeError, ValueError) as e:
                # Misconfiguration of thread settings is non-fatal
                logging.warning(f"Impossible de configurer num_threads: {e}")
                
            # Mode évaluation par défaut
            torch.set_grad_enabled(False)
            
            # Désactiver JIT et optimisations sélectivement
            if hasattr(torch, 'jit'):
                try:
                    torch._C._jit_set_profiling_mode(False)
                except:
                    pass
                try:
                    torch._C._jit_set_profiling_executor(False)
                except:
                    pass
            
            # Éviter les conflits de classes si possible
            if hasattr(torch, '_C') and hasattr(torch._C, '_disable_torch_function_mode'):
                try:
                    torch._C._disable_torch_function_mode()
                except:
                    pass
            
            _pytorch_configured = True
            logging.info("PyTorch configuré avec succès pour Streamlit")
            return True
            
        except (ImportError, RuntimeError, AttributeError) as e:
            # Import or configuration errors should disable PyTorch support
            logging.warning(f"Configuration PyTorch échouée: {e}")
            return False

# Configurer PyTorch immédiatement
PYTORCH_AVAILABLE = configure_pytorch_safe()

# === IMPORTS POUR TRAITEMENT DE DOCUMENTS ===
try:
    import pdfplumber
    from pdf2docx import parse as pdf2docx_parse
    from docx import Document
    PDF_SUPPORT = True
    logging.info("Support PDF activé")
except ImportError as e:
    PDF_SUPPORT = False
    logging.warning(f"Support PDF désactivé: {e}")

# === IMPORTS IA SÉCURISÉS ===
AI_SUPPORT = False
SPACY_SUPPORT = False
TRANSFORMERS_SUPPORT = False

# Configuration SpaCy (plus stable avec Streamlit)
try:
    import spacy
    SPACY_SUPPORT = True
    logging.info("SpaCy disponible")
except ImportError:
    logging.warning("SpaCy non disponible")

# Configuration Transformers (avec protection anti-conflit)
if PYTORCH_AVAILABLE:
    try:
        # Import sécurisé de transformers
        with _pytorch_lock:
            import transformers
            from transformers import pipeline
            
            # Réduire les logs transformers
            transformers.logging.set_verbosity_error()
            
            TRANSFORMERS_SUPPORT = True
            AI_SUPPORT = True
            logging.info("Transformers configuré avec succès")
            
    except (ImportError, OSError, RuntimeError) as e:
        # Import or runtime issues while loading transformers
        logging.warning(f"Transformers non disponible: {e}")

# === IMPORTS LOCAUX ===
# Assuming these are in a config.py file
ENTITY_PATTERNS = {
    "EMAIL": r"\b[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}\b",  # RFC 5322
    "PHONE": r"\b(?:\+33|0)(?:[1-5]|[67])(?:[ .-]?\d{2}){4}\b",  # Plan national ARCEP
    "IBAN": r"\bFR\d{2}\s?\d{5}\s?\d{5}\s?[A-Z0-9]{11}\s?\d{2}\b",  # IBAN FR ISO 13616
    "DATE": r"\b(?:0?[1-9]|[12][0-9]|3[01])[-/](?:0?[1-9]|1[0-2])[-/](?:19|20)\d{2}\b"  # JJ/MM/AAAA
}

ENTITY_COLORS = {
    "PERSON": "#FF6B6B",
    "ORG": "#4ECDC4",
    "EMAIL": "#45B7D1"
}

DEFAULT_REPLACEMENTS = {
    "PERSON": "[PERSONNE]",
    "ORG": "[ORGANISATION]",
    "EMAIL": "[EMAIL]",
    "PHONE": "[PHONE]",
    "IBAN": "[IBAN]",
    "DATE": "[DATE]"
}

# Types d'entités disposant de validations fortes
VALIDATED_ENTITY_TYPES = {"EMAIL", "PHONE", "IBAN", "SIRET", "SIREN", "SSN", "TVA"}

# === FILTRES FRANÇAIS ===
try:
    from spacy.lang.fr.stop_words import STOP_WORDS as FRENCH_STOP_WORDS
except Exception:  # pragma: no cover - fallback si SpaCy indisponible
    FRENCH_STOP_WORDS = {
        "le", "la", "les", "de", "des", "du", "un", "une",
        "et", "en", "dans", "que", "qui", "pour", "par"
    }

FRENCH_TITLES = {
    "maître", "maitre", "m.", "mr", "mme", "mlle", "dr", "me",
    "docteur", "professeur"
}

MIN_ENTITY_LENGTH = {
    "PERSON": 2,
    "ORG": 2,
    "LOC": 2,
}

DEFAULT_FILTER_CONFIG = {
    "stopwords": True,
    "min_length": True,
    "capitalization": True,
    "title_check": True,
    "require_title": False,
}


PERSON_TITLE_PATTERN = re.compile(
    r"^(?:m\.?|mme|mlle|mle|mr|dr|me|ma[iî]tre)\s+",
    re.IGNORECASE,
)


def normalize_person_name(name: str) -> str:
    """Normalize a person name by removing titles, accents and case."""
    name = PERSON_TITLE_PATTERN.sub("", name.strip())
    name = unicodedata.normalize("NFD", name)
    name = "".join(c for c in name if unicodedata.category(c) != "Mn")
    name = name.lower()
    name = re.sub(r"[^a-z\s-]", "", name)
    name = re.sub(r"\s+", " ", name).strip()
    return name


def get_preceding_token(text: str, start: int) -> str:
    """Récupérer le mot précédent une position donnée."""
    before = text[:start].rstrip()
    match = re.search(r"(\w+)$", before)
    return match.group(1) if match else ""

# === PATTERNS FRANÇAIS AMÉLIORÉS ===
FRENCH_ENTITY_PATTERNS = {
    **ENTITY_PATTERNS,

    # Noms français avec titres de civilité
    "PERSON_FR": r"\b(?:M\.?|Mme\.?|Mlle\.?|Dr\.?|Prof\.?|Me\.?|Maître)\s+(?!(?:[Ll]e|[Ll]a|[Ll]es|[Pp]résident(?:e)?s?)\b)[A-ZÀ-Ÿ][a-zà-ÿ]+(?:\s+[A-ZÀ-Ÿ][a-zà-ÿ]+)*",  # Titres de civilité

    # Organisations françaises spécifiques
    "ORG_FR": r"\b(?:SARL|SAS|SA|SNC|EURL|SASU|SCI|SELARL|SELCA|SELAS|Association|Société|Entreprise|Cabinet|Étude|Bureau|Groupe|Fondation|Institut|Centre|Établissement)\s+[A-ZÀ-Ÿ][A-Za-zÀ-ÿ\s\-'&]+",  # Formes juridiques

    # Numéros français spécialisés
    "SSN_FR": r"\b[12]\d{2}(?:0[1-9]|1[0-2])(?:[0-9]{2}|2A|2B)\d{3}\d{3}\d{2}\b",  # NIR 15 chiffres
    "SIRET_FR": r"\b\d{3}\s?\d{3}\s?\d{3}\s?\d{5}\b",  # SIRET 14 chiffres
    "SIREN_FR": r"\b\d{3}\s?\d{3}\s?\d{3}\b",  # SIREN 9 chiffres
    "TVA_FR": r"\bFR\d{2}\d{9}\b",  # TVA intracommunautaire

    # Adresses et téléphones français
    "FRENCH_ADDRESS": r"\b\d{1,4}\s?(?:bis|ter|quater)?\s+(?:rue|avenue|boulevard|place|square|impasse|allée|chemin|route|passage|villa|cité|quai|esplanade|parvis|cours|mail|faubourg)\s+[A-Za-zÀ-ÿ'\-\s]+,?\s\d{5}\s[A-Za-zÀ-ÿ'\-\s]+\b",  # Adresse postale
    "FRENCH_MOBILE": r"\b(?:\+33|0)[67](?:[ .-]?\d{2}){4}\b",  # Mobile 06/07
    "FRENCH_LANDLINE": r"\b(?:\+33|0)[1-5](?:[ .-]?\d{2}){4}\b",  # Fixe 01-05

    # Références juridiques
    "ARTICLE_LOI": r"\b[Aa]rt(?:icle)?\.?\s+[LDR]\d+(?:-\d+)*\s+du\s+Code\s+[A-Za-zÀ-ÿ\s]+\b",  # Article de loi
    "NUMERO_DOSSIER": r"\b(?:n°|N°|numéro|Numéro)\s*:?\s*\d{2,}/\d{2,}(?:/\d{2,})?\b",  # Numéro de dossier
    "RG_NUMBER": r"\bRG\s*:?\s*\d{2}/\d{5}\b"  # Référence de greffe
}

# === MODÈLES IA OPTIMISÉS ===
AI_MODELS = {
    "french_spacy_lg": {
        "name": "fr_core_news_lg",
        "type": "spacy",
        "language": "fr",
        "description": "Modèle SpaCy français large (recommandé)"
    },
    "french_spacy_sm": {
        "name": "fr_core_news_sm", 
        "type": "spacy",
        "language": "fr",
        "description": "Modèle SpaCy français compact"
    },
    "french_camembert": {
        "name": "Jean-Baptiste/camembert-ner",
        "type": "transformers",
        "language": "fr",
        "description": "CamemBERT NER français"
    },
    "multilingual_bert": {
        "name": "dbmdz/bert-large-cased-finetuned-conll03-english",
        "type": "transformers",
        "language": "multi",
        "description": "BERT multilingue"
    },
    "distilbert_lightweight": {
        "name": "distilbert-base-cased",
        "type": "transformers",
        "language": "en",
        "description": "DistilBERT léger (fallback)"
    }
}

@dataclass
class Entity:
    """Classe représentant une entité détectée avec informations enrichies"""
    id: str
    type: str
    value: str
    start: int
    end: int
    confidence: float = 1.0
    replacement: Optional[str] = None
    page: Optional[int] = None
    context: Optional[str] = None
    method: Optional[str] = "regex"  # "regex", "spacy", "transformers"
    source_model: Optional[str] = None
    # Métadonnées de déduplication
    total_occurrences: int = 1
    variants: Optional[List[str]] = None
    all_positions: Optional[List[Tuple[int, int]]] = None

class RegexAnonymizer:
    """Anonymiseur Regex avancé avec patterns français optimisés"""

    def __init__(
        self,
        use_french_patterns: bool = True,
        *,
        algorithm: str = "rapidfuzz",
        score_cutoff: Optional[float] = None,
        titles: Optional[List[str]] = None,
    ):
        self.patterns = FRENCH_ENTITY_PATTERNS if use_french_patterns else ENTITY_PATTERNS
        self.replacements = DEFAULT_REPLACEMENTS
        self.use_french_patterns = use_french_patterns
        # Mapping des valeurs d'entités vers leurs jetons anonymisés
        self.entity_mapping: Dict[str, Dict[str, Dict[str, Any]]] = {}
        # Compteurs de jetons générés par type d'entité
        self.entity_counters: Dict[str, int] = {}
        # Historique des fusions de valeurs vers un même jeton
        self.merge_history: List[Dict[str, Any]] = []
        # Patterns pour filtrer les faux positifs
        self.false_positive_patterns = self._load_false_positive_patterns()
        # Paramètres de similarité
        self.algorithm = algorithm
        self.score_cutoff = (
            score_cutoff if score_cutoff is not None else get_similarity_threshold()
        )
        self.similarity_weights = get_similarity_weights()
        # Pré-chargement de la fonction de similarité pour éviter les reimportations
        self._similarity_func = None
        try:
            if self.algorithm == "rapidfuzz":
                from rapidfuzz import fuzz

                self._similarity_func = lambda a, b: fuzz.ratio(a, b) / 100.0
            elif self.algorithm in {"levenshtein", "python-Levenshtein"}:
                from Levenshtein import ratio  # type: ignore

                self._similarity_func = ratio
            else:  # pragma: no cover - unexpected algorithm
                raise ImportError
        except ImportError:
            logging.warning(
                "Similarity algorithm %s not available, falling back to Levenshtein",
                self.algorithm,
            )
            self.algorithm = "levenshtein"
            try:
                from Levenshtein import ratio  # type: ignore

                self._similarity_func = ratio
            except ImportError:
                self._similarity_func = None
        # Titres utilisés pour la normalisation des noms
        self.titles = titles if titles is not None else get_name_normalization_titles()
        # Normaliseur légal pour les noms de personnes
        self.legal_normalizer = LegalEntityNormalizer(
            titles=set(self.titles) | LegalEntityNormalizer.DEFAULT_TITLES,
            score_cutoff=self.score_cutoff,
            weights=self.similarity_weights,
        )
        # Structures de recherche par similarité (BK-tree)
        self.bk_trees: Dict[str, BKTree] = {}
        self.bktree_threshold = 100
        # Mapping canonique -> jeton exposé aux composants externes
        self.canonical_token_map: Dict[str, Dict[str, str]] = {}
        logging.info(
            f"RegexAnonymizer initialisé avec {len(self.patterns)} patterns"
        )

    def _similarity_score(self, a: str, b: str) -> float:
        """Compute similarity score between two strings using the pre-loaded algorithm."""
        if self._similarity_func is None:
            return 0.0
        return self._similarity_func(a, b)

    def _load_terms_from_file(self, path: Path) -> List[str]:
        """Load terms from a file, one per line."""
        try:
            with open(path, "r", encoding="utf-8") as f:
                return [line.strip() for line in f if line.strip() and not line.startswith("#")]
        except OSError:
            logging.warning(f"Term file not found: {path}")
            return []

    def _load_false_positive_patterns(self) -> Dict[str, Optional[re.Pattern]]:
        """Load false positive patterns from external lists and env/config."""
        base_dir = Path(__file__).parent / "resources"
        lists = {
            "cities": self._load_terms_from_file(base_dir / "french_cities.txt"),
            "legal": self._load_terms_from_file(base_dir / "legal_terms.txt"),
            "titles": self._load_terms_from_file(base_dir / "professional_titles.txt"),
        }

        config_path = os.getenv("ANONYMIZER_EXTRA_TERMS_FILE")
        if config_path and Path(config_path).is_file():
            try:
                with open(config_path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                for key in lists:
                    lists[key].extend(data.get(key, []))
            except (OSError, json.JSONDecodeError) as e:
                logging.warning(f"Could not load extra terms file {config_path}: {e}")

        env_vars = {
            "cities": os.getenv("ANONYMIZER_EXTRA_CITIES", ""),
            "legal": os.getenv("ANONYMIZER_EXTRA_LEGAL_TERMS", ""),
            "titles": os.getenv("ANONYMIZER_EXTRA_TITLES", ""),
        }
        for key, value in env_vars.items():
            if value:
                lists[key].extend([v.strip() for v in value.split(",") if v.strip()])

        patterns: Dict[str, Optional[re.Pattern]] = {}
        for key, terms in lists.items():
            terms = [t.strip() for t in terms if t.strip()]
            if terms:
                escaped_terms = [re.escape(t) for t in terms]
                pattern = re.compile(r"\b(?:" + "|".join(escaped_terms) + r")\b", re.IGNORECASE)
                patterns[key] = pattern
            else:
                patterns[key] = None
        return patterns
    
    def detect_entities(
        self,
        text: str,
        min_confidence: float = 0.0,
        compute_conf: bool = True,
    ) -> List[Entity]:
        """Détection d'entités avec patterns regex optimisés"""
        raw_entities: List[Entity] = []
        entity_id = 0

        for entity_type, pattern in self.patterns.items():
            try:
                compiled_pattern = re.compile(pattern, re.IGNORECASE | re.MULTILINE)

                for match in compiled_pattern.finditer(text):
                    # Normaliser le type d'entité
                    normalized_type = self._normalize_entity_type(entity_type)

                    # Validation de l'entité
                    if not self._is_valid_entity_match(match.group(), normalized_type):
                        continue

                    entity = Entity(
                        id=f"regex_{entity_id}",
                        type=normalized_type,
                        value=match.group().strip(),
                        start=match.start(),
                        end=match.end(),
                        confidence=1.0,
                        replacement=self.replacements.get(normalized_type, f"[{normalized_type}]"),
                        context=self._extract_context(text, match.start(), match.end()),
                        method="regex"
                    )
                    raw_entities.append(entity)
                    entity_id += 1

            except re.error as e:
                logging.warning(f"Pattern regex invalide pour {entity_type}: {e}")
                continue

        # Déduplication directe des entités détectées
        entities = self._deduplicate_entities(raw_entities, text)

        # Post-traitement: éliminer chevauchements et nettoyer
        entities = self._remove_overlapping_entities(entities)
        entities = self._clean_entities(entities)

        if compute_conf:
            for entity in entities:
                validation_score = 1.0 if entity.type in VALIDATED_ENTITY_TYPES else 0.5
                entity.confidence = compute_confidence(1.0, validation_score, 0.0)
            entities = [e for e in entities if e.confidence >= min_confidence]

        logging.info(f"RegexAnonymizer: {len(entities)} entités détectées")
        return entities

    def anonymize_text(
        self,
        text: str,
        entities: List[Entity],
        entity_manager: Optional["EntityManager"] = None,
    ) -> Tuple[str, Dict[str, Dict[str, Dict[str, Any]]]]:
        """Anonymise le texte en remplaçant les entités détectées.

        Returns
        -------
        Tuple[str, Dict[str, Dict[str, Dict[str, Any]]]]
            Le texte anonymisé et le mapping des valeurs originales vers les
            jetons générés par type d'entité.
        """
        # Réinitialiser le mapping et les compteurs pour chaque nouveau document
        self.entity_mapping = {}
        self.entity_counters = {}
        self.canonical_token_map = {}

        # Première passe : construire le mapping des entités
        replacements: List[Tuple[str, str]] = []

        def _names_match(a: str, b: str) -> bool:
            if a == b or a in b or b in a:
                return True
            pa, pb = a.split(), b.split()
            if pa and pb and pa[-1] == pb[-1]:
                fa, fb = pa[0], pb[0]
                if (len(fa) == 1 and fb.startswith(fa)) or (
                    len(fb) == 1 and fa.startswith(fb)
                ):
                    return True
            return False

        person_entities = [e for e in entities if e.type == "PERSON"]
        other_entities = [e for e in entities if e.type != "PERSON"]

        if person_entities:
            type_map = self.entity_mapping.setdefault("PERSON", {})
            canonical_map = self.canonical_token_map.setdefault("PERSON", {})
            normalized_variants: Dict[str, set] = {}
            for ent in person_entities:
                norm = normalize_person_name(ent.value)
                normalized_variants.setdefault(norm, set()).add(ent.value)

            for norm in sorted(normalized_variants.keys(), key=len, reverse=True):
                token = None
                for existing_norm, existing_token in canonical_map.items():
                    if _names_match(norm, existing_norm):
                        token = existing_token
                        break
                if token is None:
                    count = self.entity_counters.get("PERSON", 0) + 1
                    self.entity_counters["PERSON"] = count
                    token = f"[PERSON_{count}]"
                canonical_map[norm] = token
                type_map[norm] = {
                    "token": token,
                    "variants": normalized_variants[norm],
                    "canonical": norm,
                    "origin": next(iter(normalized_variants[norm])),
                    "origin_timestamp": datetime.now().isoformat(),
                }
                for variant in normalized_variants[norm]:
                    replacements.append((variant, token))
                    if entity_manager:
                        entity_manager.update_token_variants(token, variant)

            for ent in person_entities:
                norm = normalize_person_name(ent.value)
                ent.replacement = canonical_map.get(norm)

        for entity in other_entities:
            type_map = self.entity_mapping.setdefault(entity.type, {})
            canonical_map = self.canonical_token_map.setdefault(entity.type, {})
            tree = self.bk_trees.get(entity.type)
            canonical = entity.value
            token = None
            best_score = 0.0
            best_key: Optional[str] = None
            if len(canonical_map) >= self.bktree_threshold and RFLevenshtein is not None:
                if tree is None:
                    tree = BKTree(RFLevenshtein.distance)
                    for key in canonical_map.keys():
                        tree.add(key)
                    self.bk_trees[entity.type] = tree
                max_dist = max(1, int((1 - self.score_cutoff) * len(canonical)))
                candidates = tree.search(canonical, max_dist)
                for cand, dist in candidates:
                    score = 1 - dist / max(len(canonical), len(cand))
                    if score >= self.score_cutoff and score > best_score:
                        best_score = score
                        token = canonical_map[cand]
                        best_key = cand
            else:
                for existing_key, existing_token in canonical_map.items():
                    if canonical in existing_key or existing_key in canonical:
                        score = self._similarity_score(canonical, existing_key)
                        if score >= self.score_cutoff and score > best_score:
                            best_score = score
                            token = existing_token
                            best_key = existing_key

            if token is None:
                count = self.entity_counters.get(entity.type, 0) + 1
                self.entity_counters[entity.type] = count
                token = f"[{entity.type}_{count}]"
                canonical_map[canonical] = token
                if tree is not None:
                    tree.add(canonical)
                type_map[canonical] = {
                    "token": token,
                    "variants": {entity.value},
                    "canonical": canonical,
                    "origin": entity.value,
                    "origin_timestamp": datetime.now().isoformat(),
                }
            else:
                canonical_map[canonical] = token
                if tree is not None:
                    tree.add(canonical)
                if best_key and best_key in type_map:
                    entry = type_map[best_key]
                    if entity.value not in entry["variants"]:
                        entry["variants"].add(entity.value)
                        self.merge_history.append(
                            {
                                "original": entry.get("origin", best_key),
                                "variant": entity.value,
                                "token": token,
                                "timestamp": datetime.now().isoformat(),
                            }
                        )
                    if entity_manager:
                        entity_manager.update_token_variants(token, entity.value)
                else:
                    type_map[canonical] = {
                        "token": token,
                        "variants": {entity.value},
                        "canonical": canonical,
                        "origin": entity.value,
                        "origin_timestamp": datetime.now().isoformat(),
                    }

            entity.replacement = token
            replacements.append((entity.value, token))

        # Deuxième passe : appliquer les remplacements directement sur les valeurs
        for original, token in sorted(set(replacements), key=lambda x: len(x[0]), reverse=True):
            pattern = r"\b" + re.escape(original) + r"\b"
            text = re.sub(pattern, token, text)

        return text, self.entity_mapping

    def export_merge_history(self, file_path: Optional[str] = None) -> str:
        """Exporter l'historique des fusions au format JSON.

        Parameters
        ----------
        file_path : Optional[str]
            Si fourni, le JSON est écrit dans ce fichier.

        Returns
        -------
        str
            Chaîne JSON représentant l'historique complet.
        """
        data = json.dumps(self.merge_history, ensure_ascii=False, indent=2)
        if file_path:
            try:
                with open(file_path, "w", encoding="utf-8") as f:
                    f.write(data)
            except OSError:
                logging.warning("Impossible d'écrire le fichier d'historique: %s", file_path)
        return data

    def _validate_date_fr(self, date_str: str, context: Optional[str] = None) -> bool:
        """Validation d'une date française avec vérification du contexte légal"""
        if not date_str:
            return False

        # Vérifier la validité de la date (jour/mois/année)
        try:
            sep = "/" if "/" in date_str else "-"
            parsed = datetime.strptime(date_str, f"%d{sep}%m{sep}%Y")
            if parsed.year < 1900 or parsed.year > 2100:
                return False
        except ValueError:
            return False

        # Vérifier que le contexte n'indique pas un article ou une référence légale
        if context:
            lowered = context.lower().replace("**", "")
            try:
                idx = lowered.index(date_str.lower())
            except ValueError:
                idx = -1

            if idx != -1:
                start = max(0, idx - 20)
                end = min(len(lowered), idx + len(date_str) + 20)
                window = lowered[start:end]
                if re.search(r"\b(article|art\.|loi|décret|code)\b", window):
                    return False

        return True

    def _clean_entities(self, entities: List[Entity]) -> List[Entity]:
        """Nettoyage et validation finale des entités détectées"""
        cleaned = []
        for entity in entities:
            # Vérification et nettoyage de la valeur
            if not entity.value or len(entity.value.strip()) < 2:
                continue

            # Normalisation des espaces
            entity.value = re.sub(r'\s+', ' ', entity.value.strip())

            # Validation selon le type
            if entity.type == "EMAIL":
                if not "@" in entity.value or not "." in entity.value.split("@")[1]:
                    logging.warning(f"Email invalide ignoré: {entity.value}")
                    continue
            elif entity.type == "PHONE":
                if not EntityValidator.validate_phone_fr(entity.value):
                    logging.warning(f"Téléphone invalide ignoré: {entity.value}")
                    continue
            elif entity.type == "IBAN":
                if not EntityValidator.validate_iban_fr(entity.value):
                    logging.warning(f"IBAN invalide ignoré: {entity.value}")
                    continue
            elif entity.type == "DATE":
                if not self._validate_date_fr(entity.value, entity.context):
                    logging.warning(f"Date invalide ignorée: {entity.value}")
                    continue
            elif entity.type == "SIREN":
                if not EntityValidator.validate_siren(entity.value):
                    logging.warning(f"SIREN invalide ignoré: {entity.value}")
                    continue
            elif entity.type == "SIRET":
                if not EntityValidator.validate_siret(entity.value):
                    logging.warning(f"SIRET invalide ignoré: {entity.value}")
                    continue
            elif entity.type == "SSN":
                if not EntityValidator.validate_ssn_fr(entity.value):
                    logging.warning(f"SSN invalide ignoré: {entity.value}")
                    continue

            # Ajout si valide
            cleaned.append(entity)

        return cleaned
    
    def _normalize_entity_type(self, entity_type: str) -> str:
        """Normaliser les types d'entités français"""
        mapping = {
            "PERSON_FR": "PERSON",
            "ORG_FR": "ORG", 
            "SSN_FR": "SSN",
            "SIRET_FR": "SIRET",
            "SIREN_FR": "SIREN",
            "TVA_FR": "TVA",
            "ARTICLE_LOI": "LEGAL_REF",
            "NUMERO_DOSSIER": "CASE_NUMBER",
            "RG_NUMBER": "COURT_REF"
        }
        return mapping.get(entity_type, entity_type)
    
    def _is_valid_entity_match(self, value: str, entity_type: str) -> bool:
        """Valider qu'une entité détectée est pertinente"""
        value = value.strip()
        
        # Longueur minimale
        if len(value) < 2:
            return False
        
        # Filtres spécifiques par type
        if entity_type == "EMAIL":
            valid = "@" in value and "." in value.split("@")[-1]
            if not valid:
                logging.warning(f"Match email invalide ignoré: {value}")
            return valid
        elif entity_type == "PHONE":
            valid = EntityValidator.validate_phone_fr(value)
            if not valid:
                logging.warning(f"Match téléphone invalide ignoré: {value}")
            return valid
        elif entity_type == "IBAN":
            valid = EntityValidator.validate_iban_fr(value)
            if not valid:
                logging.warning(f"Match IBAN invalide ignoré: {value}")
            return valid
        elif entity_type == "SIREN":
            valid = EntityValidator.validate_siren(value)
            if not valid:
                logging.warning(f"Match SIREN invalide ignoré: {value}")
            return valid
        elif entity_type == "SIRET":
            valid = EntityValidator.validate_siret(value)
            if not valid:
                logging.warning(f"Match SIRET invalide ignoré: {value}")
            return valid
        elif entity_type == "SSN":
            valid = EntityValidator.validate_ssn_fr(value)
            if not valid:
                logging.warning(f"Match SSN invalide ignoré: {value}")
            return valid

        return True
    
    def _extract_context(self, text: str, start: int, end: int, context_length: int = 80) -> str:
        """Extraire le contexte enrichi autour d'une entité"""
        context_start = max(0, start - context_length)
        context_end = min(len(text), end + context_length)
        
        context = text[context_start:context_end]
        entity_value = text[start:end]
        
        # Position relative dans le contexte
        relative_start = start - context_start
        relative_end = end - context_start
        
        # Marquer l'entité dans le contexte
        highlighted_context = (
            context[:relative_start] + 
            f"**{entity_value}**" + 
            context[relative_end:]
        )

        return highlighted_context.strip()

    def _get_person_signature(self, name: str) -> str:
        """Normalise un nom de personne pour la déduplication"""
        clean = re.sub(r"^(M\.?|Mme\.?|Mlle\.?|Dr\.?|Prof\.?|Me\.?|Maître)\s+", "", name, flags=re.IGNORECASE)
        return clean.lower().strip()

    def _get_entity_signature(self, entity: Entity) -> Tuple[str, str]:
        """Retourne une signature unique pour l'entité"""
        if entity.type == "PERSON":
            sig = self._get_person_signature(entity.value)
        else:
            sig = entity.value.lower().strip()
        return (entity.type, sig)

    def _count_all_occurrences(self, text: str, entity_list: List[Entity]) -> int:
        """Compte toutes les occurrences de toutes les variantes"""
        total = 0
        seen_positions: Set[Tuple[int, int]] = set()
        for ent in entity_list:
            for match in re.finditer(re.escape(ent.value), text, re.IGNORECASE):
                pos = (match.start(), match.end())
                if pos not in seen_positions:
                    seen_positions.add(pos)
                    total += 1
        return total

    def _deduplicate_entities(self, raw_entities: List[Entity], text: str) -> List[Entity]:
        """Fusionne les entités ayant la même signature"""
        groups: Dict[Tuple[str, str], List[Entity]] = {}
        for ent in raw_entities:
            signature = self._get_entity_signature(ent)
            groups.setdefault(signature, []).append(ent)

        final_entities: List[Entity] = []
        for (etype, _sig), group in groups.items():
            best_entity = max(group, key=lambda e: len(e.value))
            total_count = self._count_all_occurrences(text, group)
            final_entity = Entity(
                id=best_entity.id,
                type=etype,
                value=best_entity.value,
                start=best_entity.start,
                end=best_entity.end,
                confidence=best_entity.confidence,
                replacement=best_entity.replacement,
                page=best_entity.page,
                context=best_entity.context,
                method=best_entity.method,
                source_model=best_entity.source_model,
                total_occurrences=total_count,
                variants=list(dict.fromkeys([e.value for e in group])),
                all_positions=[(e.start, e.end) for e in group],
            )
            final_entities.append(final_entity)

        return final_entities
    
    def _remove_overlapping_entities(self, entities: List[Entity]) -> List[Entity]:
        """Éliminer les chevauchements avec priorité intelligente"""
        if not entities:
            return entities
        
        # Trier par position puis par longueur
        sorted_entities = sorted(entities, key=lambda x: (x.start, -(x.end - x.start)))
        filtered_entities = []
        
        for entity in sorted_entities:
            overlaps = False
            
            for i, accepted in enumerate(filtered_entities):
                if self._entities_overlap(entity, accepted):
                    # Résoudre le conflit selon des règles de priorité
                    winner = self._resolve_conflict(entity, accepted)

                    if winner == entity:
                        filtered_entities[i] = entity
                    
                    overlaps = True
                    break
            
            if not overlaps:
                filtered_entities.append(entity)
        
        return filtered_entities
    
    def _entities_overlap(self, entity1: Entity, entity2: Entity) -> bool:
        """Vérification de chevauchement"""
        return not (entity1.end <= entity2.start or entity2.end <= entity1.start)
    
    def _resolve_conflict(self, entity1: Entity, entity2: Entity) -> Entity:
        """Résolution de conflit avec priorités intelligentes"""
        # Priorité 1: Types structurés
        structured_types = {'EMAIL', 'PHONE', 'IBAN', 'SIRET', 'SIREN', 'SSN', 'TVA'}
        
        if entity1.type in structured_types and entity2.type not in structured_types:
            return entity1
        elif entity2.type in structured_types and entity1.type not in structured_types:
            return entity2
        
        # Priorité 2: Méthode (regex > spacy > transformers pour les conflits)
        method_priority = {'regex': 3, 'spacy': 2, 'transformers': 1}
        
        priority1 = method_priority.get(entity1.method, 0)
        priority2 = method_priority.get(entity2.method, 0)
        
        if priority1 != priority2:
            return entity1 if priority1 > priority2 else entity2
        
        # Priorité 3: Confiance
        if abs(entity1.confidence - entity2.confidence) > 0.15:
            return entity1 if entity1.confidence > entity2.confidence else entity2
        
        # Priorité 4: Longueur (plus spécifique)
        len1 = entity1.end - entity1.start
        len2 = entity2.end - entity2.start
        
        return entity1 if len1 >= len2 else entity2
    
    def _validate_anonymization(self, original_text: str, anonymized_text: str, entities: List[Entity]) -> Dict[str, Any]:
        """Validation complète de l'anonymisation"""
        validation = {
            "success": True,
            "issues": [],
            "warnings": [],
            "stats": {}
        }
        
        # Vérifier que les entités ont été remplacées
        not_replaced = []
        for entity in entities:
            if entity.value in anonymized_text:
                not_replaced.append(entity.value)
        
        if not_replaced:
            validation["success"] = False
            validation["issues"].extend([f"Non remplacé: {val}" for val in not_replaced[:5]])
        
        # Détection de fuites potentielles
        potential_leaks = self._detect_potential_leaks(anonymized_text)
        if potential_leaks:
            validation["warnings"].extend(potential_leaks)
        
        # Statistiques
        validation["stats"] = {
            "original_length": len(original_text),
            "anonymized_length": len(anonymized_text),
            "entities_replaced": len(entities),
            "not_replaced_count": len(not_replaced),
            "potential_leaks": len(potential_leaks),
            "success_rate": (len(entities) - len(not_replaced)) / len(entities) * 100 if entities else 100
        }
        
        return validation
    
    def _detect_potential_leaks(self, text: str) -> List[str]:
        """Détecter des fuites potentielles dans le texte anonymisé"""
        leaks = []
        
        # Patterns pour détecter des données manquées
        leak_patterns = {
            "email_like": r'\b[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}\b',
            "phone_like": r'\b(?:\+33|0)[1-9](?:[0-9\s.-]{8,})\b',
            "iban_like": r'\b[A-Z]{2}\d{2}[A-Z0-9]{4}\d{7}[A-Z0-9]*\b',
            "potential_name": r'\b[A-Z][a-z]{2,}\s+[A-Z][a-z]{2,}\b',
            "date_like": r'\b\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4}\b'
        }
        
        for pattern_name, pattern in leak_patterns.items():
            matches = re.findall(pattern, text)
            if matches:
                # Filtrer les faux positifs connus
                filtered_matches = self._filter_false_positives(matches, pattern_name)
                if filtered_matches:
                    leaks.append(f"Fuite potentielle {pattern_name}: {len(filtered_matches)} occurrence(s)")
        
        return leaks
    
    def _filter_false_positives(self, matches: List[str], pattern_type: str) -> List[str]:
        """Filtrer les faux positifs selon le type"""
        if pattern_type == "potential_name":
            patterns = self.false_positive_patterns
            filtered = []
            for m in matches:
                if any(p and p.search(m) for p in patterns.values()):
                    continue
                filtered.append(m)
            return filtered

        elif pattern_type == "phone_like":
            # Vérifier que c'est vraiment un numéro de téléphone
            return [m for m in matches if len(re.sub(r'\D', '', m)) >= 8]

        return matches
    
    def _generate_processing_stats(
        self,
        entities: List[Entity],
        text: str,
        processing_time: float,
        thresholds: Optional[Dict[str, float]] = None,
    ) -> Dict[str, Any]:
        """Générer des statistiques complètes de traitement"""
        entity_types = {}
        confidence_values = []
        method_counts = {}
        
        for entity in entities:
            # Types
            entity_type = entity.type
            entity_types[entity_type] = entity_types.get(entity_type, 0) + 1
            
            # Confiance
            if hasattr(entity, 'confidence') and entity.confidence is not None:
                confidence_values.append(entity.confidence)
            
            # Méthodes
            method = getattr(entity, 'method', 'unknown')
            method_counts[method] = method_counts.get(method, 0) + 1
        
        # Statistiques de confiance
        confidence_stats = {}
        thresholds = thresholds or {"high": 0.8, "medium": 0.5}
        if confidence_values:
            high = thresholds.get("high", 0.8)
            medium = thresholds.get("medium", 0.5)
            confidence_stats = {
                "min": min(confidence_values),
                "max": max(confidence_values),
                "average": sum(confidence_values) / len(confidence_values),
                "std": self._calculate_std(confidence_values),
                "high_confidence_count": len([c for c in confidence_values if c >= high]),
                "medium_confidence_count": len([c for c in confidence_values if medium <= c < high]),
                "low_confidence_count": len([c for c in confidence_values if c < medium])
            }
        
        return {
            "total_entities": len(entities),
            "entity_types": entity_types,
            "method_distribution": method_counts,
            "confidence_stats": confidence_stats,
            "processing_time": processing_time,
            "entities_per_second": len(entities) / processing_time if processing_time > 0 else 0,
            "text_length": len(text),
            "most_common_type": max(entity_types, key=entity_types.get) if entity_types else None,
            "ai_used": any(method in ['spacy', 'transformers'] for method in method_counts.keys())
        }
    
    def _calculate_std(self, values: List[float]) -> float:
        """Calculer l'écart-type"""
        if len(values) < 2:
            return 0.0
        
        mean = sum(values) / len(values)
        variance = sum((x - mean) ** 2 for x in values) / (len(values) - 1)
        return variance ** 0.5
    
    def _create_anonymized_document(self, original_path: str, anonymized_text: str, 
                                  metadata: Dict, entities: List[Entity]) -> str:
        """Créer le document anonymisé avec métadonnées enrichies"""
        try:
            doc = Document()
            
            # En-tête avec informations de traitement
            header = doc.sections[0].header
            header_para = header.paragraphs[0]
            header_para.text = "DOCUMENT ANONYMISÉ - CONFORME RGPD"
            
            # Page de titre
            doc.add_heading('DOCUMENT ANONYMISÉ', 0)
            
            # Informations d'anonymisation
            doc.add_heading('INFORMATIONS D\'ANONYMISATION', 1)
            
            info_para = doc.add_paragraph()
            info_para.add_run("Document original: ").bold = True
            info_para.add_run(f"{Path(original_path).name}\n")
            info_para.add_run("Date d'anonymisation: ").bold = True
            info_para.add_run(f"{datetime.now().strftime('%d/%m/%Y à %H:%M:%S')}\n")
            info_para.add_run("Méthode de détection: ").bold = True
            info_para.add_run(f"{metadata.get('detection_method', 'regex').upper()}\n")
            info_para.add_run("Entités anonymisées: ").bold = True
            info_para.add_run(f"{len(entities)}\n")
            info_para.add_run("Temps de traitement: ").bold = True
            info_para.add_run(f"{metadata.get('processing_time', 0):.2f} secondes\n")
            
            # IA utilisée si applicable
            if metadata.get('ai_used', False):
                info_para.add_run("Intelligence Artificielle: ").bold = True
                info_para.add_run("Activée (NER)\n")
            
            # Tableau récapitulatif des entités
            if entities:
                doc.add_heading('RÉCAPITULATIF DES ENTITÉS', 2)
                
                # Grouper par type
                entity_types = {}
                for entity in entities:
                    entity_type = entity.type
                    entity_types[entity_type] = entity_types.get(entity_type, 0) + 1
                
                # Créer tableau
                table = doc.add_table(rows=len(entity_types) + 1, cols=3)
                table.style = 'Table Grid'
                
                # En-têtes
                header_cells = table.rows[0].cells
                header_cells[0].text = 'Type d\'entité'
                header_cells[1].text = 'Nombre'
                header_cells[2].text = 'Pourcentage'
                
                # Données
                for i, (entity_type, count) in enumerate(sorted(entity_types.items()), 1):
                    row_cells = table.rows[i].cells
                    row_cells[0].text = entity_type
                    row_cells[1].text = str(count)
                    row_cells[2].text = f"{(count / len(entities) * 100):.1f}%"
            
            # Contenu anonymisé
            doc.add_page_break()
            doc.add_heading('CONTENU ANONYMISÉ', 1)
            
            # Diviser le texte en paragraphes
            paragraphs = anonymized_text.split('\n')
            for paragraph in paragraphs:
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())
            
            # Rapport d'audit détaillé
            doc.add_page_break()
            doc.add_heading('RAPPORT D\'AUDIT RGPD', 1)
            
            audit_para = doc.add_paragraph()
            audit_para.add_run("Conformité RGPD:\n").bold = True
            audit_para.add_run("• Article 4(5) - Pseudonymisation: Appliquée\n")
            audit_para.add_run("• Article 25 - Protection dès la conception: Respectée\n")
            audit_para.add_run("• Article 32 - Sécurité du traitement: Mise en œuvre\n")
            audit_para.add_run("• Recital 26 - Anonymisation: Conforme\n\n")
            
            audit_para.add_run("Recommandations:\n").bold = True
            audit_para.add_run("• Vérifiez manuellement les entités détectées\n")
            audit_para.add_run("• Conservez ce rapport pour traçabilité\n")
            audit_para.add_run("• Relecture recommandée avant diffusion\n")
            
            # Signature numérique et traçabilité
            footer = doc.sections[0].footer
            footer_para = footer.paragraphs[0]
            
            # Générer signature
            document_hash = hashlib.md5(anonymized_text.encode('utf-8')).hexdigest()[:16]
            document_id = str(uuid.uuid4())[:8].upper()
            
            footer_para.text = (
                f"Anonymiseur v2.0 - ID: {document_id} - "
                f"Hash: {document_hash} - "
                f"Généré: {datetime.now().strftime('%d/%m/%Y %H:%M')}"
            )
            
            # Sauvegarde
            output_path = os.path.join(self.temp_dir, f"anonymized_{uuid.uuid4().hex[:8]}.docx")
            doc.save(output_path)
            
            logging.info(f"Document anonymisé créé: {output_path}")
            return output_path
        
        except (OSError, ValueError) as e:
            # File system issues while saving the document
            logging.error(f"Erreur création document: {e}")
            raise RuntimeError(f"Impossible de créer le document anonymisé: {str(e)}")
    
    def get_processing_statistics(self) -> Dict[str, Any]:
        """Obtenir les statistiques globales"""
        stats = self.processing_stats.copy()
        
        # Ajouter des informations système
        stats.update({
            "ai_support": AI_SUPPORT,
            "spacy_support": SPACY_SUPPORT,
            "transformers_support": TRANSFORMERS_SUPPORT,
            "pdf_support": PDF_SUPPORT,
            "pytorch_available": PYTORCH_AVAILABLE
        })
        
        return stats
    
    def cleanup(self):
        """Nettoyer les ressources"""
        try:
            import shutil
            if os.path.exists(self.temp_dir):
                shutil.rmtree(self.temp_dir, ignore_errors=True)
            logging.info("Nettoyage terminé")
        except OSError as e:
            logging.warning(f"Erreur nettoyage: {e}")

# === CLASSES UTILITAIRES ===

class EntityValidator:
    """Validateur spécialisé pour les entités françaises"""
    
    @staticmethod
    def validate_email(email: str) -> bool:
        """Validation email française"""
        pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
        return bool(re.match(pattern, email))
    
    @staticmethod
    def validate_phone_fr(phone: str) -> bool:
        """Validation téléphone français"""
        digits = re.sub(r'\D', '', phone)

        # Convertir les formats internationaux en format national
        if digits.startswith('0033'):
            digits = '0' + digits[4:]
        elif digits.startswith('33'):
            digits = '0' + digits[2:]

        # Doit être au format national 10 chiffres
        if len(digits) != 10 or not digits.startswith('0'):
            return False

        # Préfixes autorisés: 01-05 (fixe) et 06-07 (mobile)
        return digits[1] in '1234567'

    @staticmethod
    def validate_ssn_fr(ssn: str) -> bool:
        """Validation du NIR français (15 chiffres)"""
        value = ssn.replace(' ', '').upper()
        if not re.match(r'^[12]\d{2}(0[1-9]|1[0-2])(?:\d{2}|2A|2B)\d{3}\d{3}\d{2}$', value):
            return False

        nir = value[:-2]
        key = int(value[-2:])
        nir_numeric = nir.replace('2A', '19').replace('2B', '18')
        try:
            expected_key = 97 - (int(nir_numeric) % 97)
        except ValueError:
            return False
        return expected_key == key

    @staticmethod
    def validate_iban_fr(iban: str) -> bool:
        """Validation de l'IBAN français (mod-97)"""
        clean = re.sub(r'\s+', '', iban).upper()
        if not clean.startswith('FR') or len(clean) != 27:
            return False

        # Réarranger et convertir en nombre
        rearranged = clean[4:] + clean[:4]
        converted = ''
        for ch in rearranged:
            if ch.isdigit():
                converted += ch
            elif 'A' <= ch <= 'Z':
                converted += str(ord(ch) - 55)
            else:
                return False

        try:
            return int(converted) % 97 == 1
        except ValueError:
            return False
    
    @staticmethod
    def validate_siren(siren: str) -> bool:
        """Validation SIREN avec algorithme de Luhn"""
        digits = re.sub(r'\D', '', siren)
        
        if len(digits) != 9:
            return False
        
        # Algorithme de Luhn pour SIREN
        total = 0
        for i, digit in enumerate(digits):
            num = int(digit)
            if i % 2 == 1:
                num *= 2
                if num > 9:
                    num = num // 10 + num % 10
            total += num
        
        return total % 10 == 0
    
    @staticmethod
    def validate_siret(siret: str) -> bool:
        """Validation SIRET avec algorithme de Luhn"""
        digits = re.sub(r'\D', '', siret)
        
        if len(digits) != 14:
            return False
        
        # Vérifier le SIREN (9 premiers chiffres)
        siren = digits[:9]
        if not EntityValidator.validate_siren(siren):
            return False
        
        # Algorithme de Luhn pour SIRET complet
        total = 0
        for i, digit in enumerate(digits):
            num = int(digit)
            if i % 2 == 0:  # Position inversée par rapport au SIREN
                num *= 2
                if num > 9:
                    num = num // 10 + num % 10
            total += num
        
        return total % 10 == 0

class AIAnonymizer:
    """Anonymiseur IA avec NER multi-modèles et gestion des conflits Streamlit"""

    def __init__(
        self,
        model_config: dict = None,
        prefer_french: bool = True,
        filter_config: Optional[Dict[str, bool]] = None,
    ):
        self.model_config = model_config or self._get_best_model(prefer_french)
        self.nlp_pipeline = None
        self.spacy_nlp = None
        self.regex_anonymizer = RegexAnonymizer(use_french_patterns=True)
        self.prefer_french = prefer_french
        self.model_loaded = False

        # Configuration des filtres
        self.filter_config = DEFAULT_FILTER_CONFIG.copy()
        if filter_config:
            self.filter_config.update(filter_config)

        # Initialiser le modèle en mode thread-safe
        self._initialize_model_safe()
    
    def _get_best_model(self, prefer_french: bool) -> dict:
        """Sélectionner le meilleur modèle disponible selon les préférences"""
        if prefer_french:
            if SPACY_SUPPORT:
                return AI_MODELS["french_spacy_lg"]
            elif TRANSFORMERS_SUPPORT:
                return AI_MODELS["french_camembert"]
            else:
                return AI_MODELS["multilingual_bert"]
        else:
            return AI_MODELS["distilbert_lightweight"]
    
    def _initialize_model_safe(self):
        """Initialisation thread-safe du modèle IA"""
        with _pytorch_lock:
            try:
                if self.model_config["type"] == "spacy":
                    self._initialize_spacy()
                else:
                    self._initialize_transformers()
                self.model_loaded = True
                
            except (OSError, RuntimeError, ValueError) as e:
                # Model loading issues disable AI features but allow regex mode
                logging.error(f"Échec du chargement du modèle IA: {e}")
                logging.info("Fallback vers mode regex uniquement")
    
    def _initialize_spacy(self):
        """Initialiser SpaCy (recommandé pour le français)"""
        if not SPACY_SUPPORT:
            raise Exception("SpaCy non disponible")
        
        try:
            self.spacy_nlp = spacy.load(self.model_config["name"])
            logging.info(f"Modèle SpaCy chargé: {self.model_config['name']}")
            
        except OSError:
            # Essayer le modèle compact si le large n'est pas disponible
            if self.model_config["name"] == "fr_core_news_lg":
                try:
                    self.spacy_nlp = spacy.load("fr_core_news_sm")
                    logging.info("Modèle SpaCy compact chargé: fr_core_news_sm")
                    return
                except OSError:
                    pass
            
            raise Exception(f"Modèle SpaCy non trouvé: {self.model_config['name']}")
        
        except (RuntimeError, ValueError) as e:
            # Other SpaCy runtime issues
            raise Exception(f"Erreur SpaCy: {e}")
    
    def _initialize_transformers(self):
        """Initialiser Transformers avec protection anti-conflit"""
        if not TRANSFORMERS_SUPPORT:
            raise Exception("Transformers non disponible")
        
        try:
            # Configuration pipeline avec protection thread
            self.nlp_pipeline = pipeline(
                "ner",
                model=self.model_config["name"],
                aggregation_strategy="simple",
                device=-1,  # Forcer CPU
                return_all_scores=False
            )
            
            logging.info(f"Pipeline Transformers chargé: {self.model_config['name']}")
            
        except (OSError, RuntimeError, ValueError) as e:
            # Fallback vers un modèle plus léger
            try:
                self.nlp_pipeline = pipeline(
                    "ner",
                    model=AI_MODELS["distilbert_lightweight"]["name"],
                    aggregation_strategy="simple",
                    device=-1
                )
                logging.info("Modèle fallback Transformers chargé")
                
            except (OSError, RuntimeError, ValueError) as e2:
                raise Exception(f"Tous les modèles Transformers ont échoué: {e}, {e2}")
    
    def detect_entities_ai(
        self,
        text: str,
        confidence_threshold: float = 0.7,
        final_threshold: float = 0.0,
    ) -> List[Entity]:
        """Détection d'entités avec IA + fusion regex"""
        ai_entities: List[Entity] = []

        # Étape 1: Détection IA
        if self.model_loaded:
            try:
                if self.spacy_nlp:
                    ai_entities = self._detect_with_spacy(text, confidence_threshold)
                elif self.nlp_pipeline:
                    ai_entities = self._detect_with_transformers(text, confidence_threshold)
                logging.info(f"IA: {len(ai_entities)} entités détectées")
            except (RuntimeError, ValueError) as e:
                logging.error(f"Erreur détection IA: {e}")

        # Étape 2: Compléter avec regex pour les entités structurées
        regex_entities = self.regex_anonymizer.detect_entities(
            text, compute_conf=False
        )
        self._compute_agreement_scores(regex_entities, ai_entities)
        merged_regex = self._merge_regex_entities(ai_entities, regex_entities)
        entities = ai_entities + merged_regex

        # Étape 3: Post-traitement final
        entities = self._post_process_entities(entities, text)
        self._apply_final_confidence(entities)
        entities = [e for e in entities if e.confidence >= final_threshold]

        logging.info(f"Total final: {len(entities)} entités")
        return entities
    
    def _detect_with_spacy(self, text: str, confidence_threshold: float) -> List[Entity]:
        """Détection avec SpaCy optimisée"""
        entities = []
        
        try:
            # Traitement par chunks pour les gros documents
            chunks = self._chunk_text(text, max_length=1000000)  # 1M chars max par chunk
            
            for chunk_start, chunk_text in chunks:
                doc = self.spacy_nlp(chunk_text)
                
                for ent in doc.ents:
                    # Calculer confiance approximative pour SpaCy
                    confidence = self._calculate_spacy_confidence(ent)
                    
                    if confidence >= confidence_threshold:
                        entity_type = self._map_spacy_label(ent.label_)
                        
                        entity = Entity(
                            id=f"spacy_{uuid.uuid4().hex[:8]}",
                            type=entity_type,
                            value=ent.text.strip(),
                            start=chunk_start + ent.start_char,
                            end=chunk_start + ent.end_char,
                            confidence=confidence,
                            replacement=DEFAULT_REPLACEMENTS.get(entity_type, f"[{entity_type}]"),
                            context=self._extract_context(text, chunk_start + ent.start_char, chunk_start + ent.end_char),
                            method="spacy",
                            source_model=self.model_config["name"]
                        )
                        entities.append(entity)
            
            return entities
            
        except (RuntimeError, ValueError) as e:
            logging.error(f"Erreur SpaCy: {e}")
            return []
    
    def _detect_with_transformers(self, text: str, confidence_threshold: float) -> List[Entity]:
        """Détection avec Transformers en mode sécurisé"""
        entities = []
        
        try:
            # Protection thread pour Transformers
            with _pytorch_lock:
                chunks = self._chunk_text(text, max_length=512)
                
                for chunk_start, chunk_text in chunks:
                    try:
                        # Pipeline NER avec gestion d'erreurs
                        ner_results = self.nlp_pipeline(chunk_text)
                        
                        for result in ner_results:
                            if result['score'] >= confidence_threshold:
                                entity_type = self._map_ner_label(result['entity_group'])
                                
                                entity = Entity(
                                    id=f"transformers_{uuid.uuid4().hex[:8]}",
                                    type=entity_type,
                                    value=result['word'].strip(),
                                    start=chunk_start + result['start'],
                                    end=chunk_start + result['end'],
                                    confidence=result['score'],
                                    replacement=DEFAULT_REPLACEMENTS.get(entity_type, f"[{entity_type}]"),
                                    context=self._extract_context(text, chunk_start + result['start'], chunk_start + result['end']),
                                    method="transformers",
                                    source_model=self.model_config["name"]
                                )
                                entities.append(entity)
                    
                    except (RuntimeError, ValueError) as e:
                        logging.warning(f"Erreur sur chunk: {e}")
                        continue
            
            return entities
            
        except (RuntimeError, ValueError) as e:
            logging.error(f"Erreur Transformers: {e}")
            return []
    
    def _chunk_text(self, text: str, max_length: int = 512) -> List[Tuple[int, str]]:
        """Diviser le texte en chunks intelligents"""
        if len(text) <= max_length:
            return [(0, text)]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + max_length
            
            # Essayer de couper à une limite de phrase ou de paragraphe
            if end < len(text):
                # Chercher un point suivi d'espace et majuscule
                sentence_end = text.rfind('. ', start, end)
                if sentence_end > start + max_length // 2:
                    end = sentence_end + 1
                else:
                    # Sinon, chercher un espace
                    space_pos = text.rfind(' ', start + max_length // 2, end)
                    if space_pos > start:
                        end = space_pos
            
            chunks.append((start, text[start:end]))
            start = end
        
        return chunks
    
    def _calculate_spacy_confidence(self, ent) -> float:
        """Calculer une confiance approximative pour SpaCy"""
        # Facteurs de confiance basés sur le type et les caractéristiques
        type_confidence = {
            'PER': 0.92, 'PERSON': 0.92,
            'ORG': 0.88,
            'LOC': 0.85, 'GPE': 0.85,
            'MISC': 0.75,
            'DATE': 0.90,
            'MONEY': 0.88,
            'PERCENT': 0.85
        }
        
        base_confidence = type_confidence.get(ent.label_, 0.80)
        
        # Ajustements
        # Longueur: entités plus longues = plus fiables
        length_factor = min(1.1, 1.0 + (len(ent.text) - 3) * 0.02)
        
        # Capitalisation: noms propres plus fiables
        if ent.text[0].isupper() and ent.label_ in ['PER', 'PERSON', 'ORG']:
            capitalization_factor = 1.05
        else:
            capitalization_factor = 1.0
        
        final_confidence = base_confidence * length_factor * capitalization_factor
        return min(0.98, final_confidence)
    
    def _map_spacy_label(self, spacy_label: str) -> str:
        """Mapper les labels SpaCy vers nos types"""
        mapping = {
            'PER': 'PERSON', 'PERSON': 'PERSON',
            'ORG': 'ORG',
            'LOC': 'LOCATION', 'GPE': 'LOCATION',
            'MISC': 'MISC',
            'DATE': 'DATE', 'TIME': 'DATE',
            'MONEY': 'MONEY',
            'PERCENT': 'PERCENT',
            'CARDINAL': 'NUMBER',
            'ORDINAL': 'NUMBER'
        }
        return mapping.get(spacy_label.upper(), spacy_label.upper())
    
    def _map_ner_label(self, ner_label: str) -> str:
        """Mapper les labels NER Transformers vers nos types"""
        mapping = {
            'PER': 'PERSON', 'PERSON': 'PERSON',
            'ORG': 'ORG', 'ORGANIZATION': 'ORG',
            'LOC': 'LOCATION', 'LOCATION': 'LOCATION',
            'MISC': 'MISC',
            'DATE': 'DATE', 'TIME': 'DATE'
        }
        return mapping.get(ner_label.upper(), ner_label.upper())
    
    def _merge_regex_entities(self, ai_entities: List[Entity], regex_entities: List[Entity]) -> List[Entity]:
        """Fusionner intelligemment les entités IA et regex"""
        merged = []
        
        for regex_entity in regex_entities:
            is_duplicate = False
            
            # Vérifier les chevauchements avec les entités IA
            for ai_entity in ai_entities:
                overlap = self._calculate_overlap(regex_entity, ai_entity)
                
                if overlap > 0.3:  # 30% de chevauchement
                    # Priorité aux entités structurées regex
                    if regex_entity.type in ['EMAIL', 'PHONE', 'IBAN', 'SIRET', 'SIREN', 'SSN', 'TVA']:
                        # Garder regex, supprimer IA
                        try:
                            ai_entities.remove(ai_entity)
                        except ValueError:
                            pass
                        break
                    else:
                        # Garder IA pour les entités non-structurées
                        is_duplicate = True
                        break
            
            if not is_duplicate:
                merged.append(regex_entity)
        
        return merged
    
    def _calculate_overlap(self, entity1: Entity, entity2: Entity) -> float:
        """Calculer le pourcentage de chevauchement entre deux entités"""
        if entity1.end <= entity2.start or entity2.end <= entity1.start:
            return 0.0  # Pas de chevauchement

        overlap_start = max(entity1.start, entity2.start)
        overlap_end = min(entity1.end, entity2.end)
        overlap_length = max(0, overlap_end - overlap_start)
        total_length = min(entity1.end, entity2.end) - max(entity1.start, entity2.start)

        return overlap_length / total_length if total_length > 0 else 0.0

    def _compute_agreement_scores(
        self, regex_entities: List[Entity], ai_entities: List[Entity]
    ) -> None:
        """Marquer les entités en accord entre regex et IA."""
        for ent in regex_entities + ai_entities:
            setattr(ent, "_agreement", 0.0)

        for r in regex_entities:
            for a in ai_entities:
                if self._calculate_overlap(r, a) >= 0.5:
                    r._agreement = 1.0
                    a._agreement = 1.0

    def _apply_final_confidence(self, entities: List[Entity]) -> None:
        """Calculer la confiance finale pour chaque entité."""
        for ent in entities:
            method_score = 1.0 if ent.method == "regex" else ent.confidence
            validation_score = (
                1.0 if ent.type in VALIDATED_ENTITY_TYPES else 0.5
            )
            agreement_score = getattr(ent, "_agreement", 0.0)
            ent.confidence = compute_confidence(
                method_score, validation_score, agreement_score
            )
    
    def _extract_context(self, text: str, start: int, end: int, context_length: int = 120) -> str:
        """Extraction de contexte intelligent"""
        context_start = max(0, start - context_length)
        context_end = min(len(text), end + context_length)
        
        # Essayer de couper à des limites de mots
        context = text[context_start:context_end]
        
        # Marquer l'entité
        entity_value = text[start:end]
        relative_start = start - context_start
        relative_end = end - context_start
        
        if 0 <= relative_start < len(context) and 0 <= relative_end <= len(context):
            highlighted_context = (
                context[:relative_start] + 
                f"**{entity_value}**" + 
                context[relative_end:]
            )
        else:
            highlighted_context = context
        
        return highlighted_context.strip()
    
    def _post_process_entities(self, entities: List[Entity], text: str) -> List[Entity]:
        """Post-traitement avec optimisations françaises"""
        processed = []

        for entity in entities:
            # Validation de base
            if not self._is_valid_entity(entity, text):
                continue

            # Nettoyage de la valeur
            entity.value = self._clean_entity_value(entity.value)

            # Amélioration du contexte
            if not entity.context:
                entity.context = self._extract_context(text, entity.start, entity.end)

            # Classification fine française
            entity.type = self._refine_french_classification(entity, text)

            # Filtres configurables
            cfg = self.filter_config

            if cfg.get("min_length", True):
                min_len = MIN_ENTITY_LENGTH.get(entity.type, 1)
                if len(entity.value) < min_len:
                    continue

            if cfg.get("stopwords", True) and entity.type == "PERSON":
                if entity.value.lower() in FRENCH_STOP_WORDS:
                    continue

            if cfg.get("capitalization", True) and entity.type == "PERSON":
                if not entity.value[:1].isupper():
                    continue

            if cfg.get("title_check", True) and entity.type == "PERSON":
                preceding = get_preceding_token(text, entity.start)
                if preceding.lower() in FRENCH_TITLES:
                    pass
                elif cfg.get("require_title", False):
                    continue

            processed.append(entity)

        # Résolution des conflits
        processed = self._resolve_entity_conflicts(processed)

        return processed
    
    def _is_valid_entity(self, entity: Entity, text: str) -> bool:
        """Validation robuste des entités"""
        # Vérifications de base
        if entity.start < 0 or entity.end > len(text):
            return False
        
        if entity.start >= entity.end:
            return False
        
        if len(entity.value.strip()) < 2:
            return False
        
        # Vérification cohérence texte
        actual_value = text[entity.start:entity.end]
        if actual_value.strip().lower() != entity.value.strip().lower():
            # Tentative de correction
            search_start = max(0, entity.start - 10)
            search_end = min(len(text), entity.end + 10)
            search_area = text[search_start:search_end]
            
            if entity.value in search_area:
                corrected_start = search_area.find(entity.value) + search_start
                entity.start = corrected_start
                entity.end = corrected_start + len(entity.value)
            else:
                return False
        
        return True
    
    def _clean_entity_value(self, value: str) -> str:
        """Nettoyage avancé des valeurs d'entités"""
        # Supprimer espaces début/fin
        cleaned = value.strip()
        
        # Supprimer ponctuation finale
        cleaned = re.sub(r'[.,;:!?]+$', '', cleaned)
        
        # Supprimer caractères bizarres
        cleaned = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', cleaned)
        
        # Normaliser espaces internes
        cleaned = re.sub(r'\s+', ' ', cleaned)
        
        return cleaned
    
    def _refine_french_classification(self, entity: Entity, text: str) -> str:
        """Classification fine pour le contexte français"""
        value = entity.value.lower()
        
        # Amélioration PERSON
        if entity.type == 'PERSON':
            context_before = text[max(0, entity.start - 30):entity.start].lower()
            
            # Vérifier titres français
            french_titles = ['monsieur', 'madame', 'mademoiselle', 'docteur', 'professeur', 'maître']
            if any(title in context_before for title in french_titles):
                return 'PERSON'
            
            # Détecter organisations
            org_words = ['société', 'entreprise', 'sarl', 'sas', 'cabinet', 'étude', 'bureau']
            if any(word in value for word in org_words):
                return 'ORG'
        
        # Amélioration ORG
        elif entity.type == 'ORG':
            legal_forms = ['sarl', 'sas', 'sa', 'snc', 'eurl', 'sasu', 'sci', 'selarl', 'selas', 'selca']
            if any(form in value for form in legal_forms):
                return 'ORG'
        
        return entity.type
    
    def _resolve_entity_conflicts(self, entities: List[Entity]) -> List[Entity]:
        """Résolution intelligente des conflits"""
        if not entities:
            return entities
        
        # Trier par position
        sorted_entities = sorted(entities, key=lambda x: x.start)
        resolved = []
        
        for current in sorted_entities:
            conflict_found = False

            for i, existing in enumerate(resolved):
                if self._entities_overlap(current, existing):
                    winner = self._resolve_conflict(current, existing)
                    resolved[i] = winner
                    conflict_found = True
                    break
            
            if not conflict_found:
                resolved.append(current)
        
        return resolved
    
    def _entities_overlap(self, entity1: Entity, entity2: Entity) -> bool:
        """Vérification de chevauchement"""
        return not (entity1.end <= entity2.start or entity2.end <= entity1.start)
    
    def _resolve_conflict(self, entity1: Entity, entity2: Entity) -> Entity:
        """Résolution de conflit avec priorités intelligentes"""
        # Priorité 1: Types structurés
        structured_types = {'EMAIL', 'PHONE', 'IBAN', 'SIRET', 'SIREN', 'SSN', 'TVA'}
        
        if entity1.type in structured_types and entity2.type not in structured_types:
            return entity1
        elif entity2.type in structured_types and entity1.type not in structured_types:
            return entity2
        
        # Priorité 2: Méthode (regex > spacy > transformers pour les conflits)
        method_priority = {'regex': 3, 'spacy': 2, 'transformers': 1}
        
        priority1 = method_priority.get(entity1.method, 0)
        priority2 = method_priority.get(entity2.method, 0)
        
        if priority1 != priority2:
            return entity1 if priority1 > priority2 else entity2
        
        # Priorité 3: Confiance
        if abs(entity1.confidence - entity2.confidence) > 0.15:
            return entity1 if entity1.confidence > entity2.confidence else entity2
        
        # Priorité 4: Longueur (plus spécifique)
        len1 = entity1.end - entity1.start
        len2 = entity2.end - entity2.start
        
        return entity1 if len1 >= len2 else entity2

class DocumentProcessor:
    """Processeur de documents avec gestion d'erreurs robuste"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc', '.txt']
        self.max_text_length = 10_000_000  # 10M caractères max
    
    def extract_text_from_pdf(self, file_path: str) -> Tuple[str, Dict]:
        """Extraction PDF robuste avec fallbacks multiples"""
        if not PDF_SUPPORT:
            raise Exception("Support PDF non disponible. Installez pdfplumber et pdf2docx.")
        
        text_content = ""
        metadata = {"pages": 0, "format": "pdf", "extraction_method": None}
        
        # Méthode 1: pdfplumber (recommandée)
        try:
            with pdfplumber.open(file_path) as pdf:
                metadata["pages"] = len(pdf.pages)
                metadata["extraction_method"] = "pdfplumber"
                
                for page_num, page in enumerate(pdf.pages, 1):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_content += f"\n--- Page {page_num} ---\n{page_text}"
                        
                        # Extraire les tableaux si peu de texte
                        if len(page_text or "") < 100:
                            tables = page.extract_tables()
                            if tables:
                                text_content += f"\n--- Tableaux Page {page_num} ---\n"
                                for table in tables:
                                    for row in table:
                                        if row:
                                            text_content += " | ".join([cell or "" for cell in row]) + "\n"
                    
                    except (RuntimeError, ValueError) as e:
                        # Skip problematic pages but continue processing others
                        logging.warning(f"Erreur page {page_num}: {e}")
                        continue
            
            if text_content.strip():
                metadata["text_length"] = len(text_content)
                return text_content, metadata
        
        except (OSError, RuntimeError) as e:
            logging.warning(f"pdfplumber échoué: {e}")
        
        # Méthode 2: PyMuPDF fallback
        try:
            import fitz
            metadata["extraction_method"] = "pymupdf"
            
            doc = fitz.open(file_path)
            metadata["pages"] = len(doc)
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                page_text = page.get_text()
                if page_text:
                    text_content += f"\n--- Page {page_num + 1} ---\n{page_text}"
            
            doc.close()
            metadata["text_length"] = len(text_content)
            return text_content, metadata
        
        except ImportError:
            logging.warning("PyMuPDF non disponible")
        except (OSError, RuntimeError, ValueError) as e:
            logging.warning(f"PyMuPDF échoué: {e}")
        
        # Méthode 3: pdf2docx puis extraction DOCX
        try:
            temp_docx = tempfile.mktemp(suffix='.docx')
            pdf2docx_parse(file_path, temp_docx)
            
            text_content, docx_metadata = self.extract_text_from_docx(temp_docx)
            metadata.update(docx_metadata)
            metadata["extraction_method"] = "pdf2docx"
            
            os.unlink(temp_docx)
            return text_content, metadata
        
        except (OSError, RuntimeError) as e:
            logging.warning(f"pdf2docx échoué: {e}")
        
        raise Exception("Impossible d'extraire le texte du PDF avec toutes les méthodes disponibles")
    
    def extract_text_from_docx(self, file_path: str) -> Tuple[str, Dict]:
        """Extraction DOCX complète avec métadonnées"""
        try:
            doc = Document(file_path)
            
            text_content = ""
            metadata = {
                "paragraphs": 0,
                "tables": 0, 
                "headers": 0,
                "footers": 0,
                "format": "docx"
            }
            
            # Extraction paragraphes
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content += para.text + "\n"
                    metadata["paragraphs"] += 1
            
            # Extraction tableaux
            for table in doc.tables:
                metadata["tables"] += 1
                text_content += "\n--- Tableau ---\n"
                
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    
                    if row_text:
                        text_content += " | ".join(row_text) + "\n"
            
            # Extraction en-têtes et pieds de page
            for section in doc.sections:
                if section.header:
                    for para in section.header.paragraphs:
                        if para.text.strip():
                            text_content += f"[EN-TÊTE] {para.text}\n"
                            metadata["headers"] += 1
                
                if section.footer:
                    for para in section.footer.paragraphs:
                        if para.text.strip():
                            text_content += f"[PIED DE PAGE] {para.text}\n"
                            metadata["footers"] += 1
            
            # Limitation de taille
            if len(text_content) > self.max_text_length:
                text_content = text_content[:self.max_text_length]
                logging.warning(f"Texte tronqué à {self.max_text_length} caractères")
            
            metadata["text_length"] = len(text_content)
            return text_content, metadata
        
        except (OSError, ValueError) as e:
            raise Exception(f"Échec extraction DOCX: {str(e)}")
    
    def extract_text_from_txt(self, file_path: str) -> Tuple[str, Dict]:
        """Extraction fichier texte avec encodage intelligent"""
        encodings = ['utf-8', 'utf-8-sig', 'latin1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text_content = f.read()
                
                # Limitation de taille
                if len(text_content) > self.max_text_length:
                    text_content = text_content[:self.max_text_length]
                    logging.warning(f"Texte tronqué à {self.max_text_length} caractères")
                
                metadata = {
                    "format": "txt",
                    "encoding": encoding,
                    "text_length": len(text_content),
                    "lines": text_content.count('\n') + 1
                }
                
                return text_content, metadata
            
            except UnicodeDecodeError:
                continue
            except OSError as e:
                raise Exception(f"Erreur lecture fichier texte: {e}")
        
        raise Exception("Impossible de décoder le fichier texte avec les encodages supportés")
    
    def process_file(self, file_path: str) -> Tuple[str, Dict]:
        """Traitement unifié avec détection automatique"""
        file_path = Path(file_path)
        file_ext = file_path.suffix.lower()
        
        if file_ext == '.pdf':
            return self.extract_text_from_pdf(str(file_path))
        elif file_ext in ['.docx', '.doc']:
            return self.extract_text_from_docx(str(file_path))
        elif file_ext == '.txt':
            return self.extract_text_from_txt(str(file_path))
        else:
            raise Exception(f"Format non supporté: {file_ext}")

class DocumentAnonymizer:
    """Anonymiseur principal avec IA fonctionnelle et optimisations Streamlit"""

    def __init__(
        self,
        prefer_french: bool = True,
        use_spacy: bool = True,
        filter_config: Optional[Dict[str, bool]] = None,
    ):
        self.regex_anonymizer = RegexAnonymizer(use_french_patterns=True)

        # Configuration des filtres
        self.filter_config = DEFAULT_FILTER_CONFIG.copy()
        if filter_config:
            self.filter_config.update(filter_config)

        # Initialisation IA conditionnelle
        if AI_SUPPORT or SPACY_SUPPORT:
            try:
                self.ai_anonymizer = AIAnonymizer(
                    prefer_french=prefer_french, filter_config=self.filter_config
                )
                logging.info("AIAnonymizer initialisé avec succès")
            except (RuntimeError, OSError, ValueError) as e:
                self.ai_anonymizer = None
                logging.warning(f"AIAnonymizer non disponible: {e}")
        else:
            self.ai_anonymizer = None
            logging.info("Mode regex uniquement (IA non disponible)")

        self.document_processor = DocumentProcessor()
        self.temp_dir = tempfile.mkdtemp()
        self.prefer_french = prefer_french
        # Mapping des entités pour assurer la cohérence des remplacements
        self.entity_mapping: Dict[str, Dict[str, Dict[str, Any]]] = {}
        # Compteurs accessibles des jetons générés par type d'entité
        self.entity_counters: Dict[str, int] = {}

        # Statistiques de traitement
        self.processing_stats = {
            "documents_processed": 0,
            "entities_detected": 0,
            "processing_time": 0.0,
            "ai_available": self.ai_anonymizer is not None
        }

    def _validate_anonymization(
        self,
        original_text: str,
        anonymized_text: str,
        entities: List[Entity],
    ) -> Dict[str, Any]:
        """Deleguer la validation à l'anonymiseur regex."""
        return self.regex_anonymizer._validate_anonymization(
            original_text, anonymized_text, entities
        )

    def process_document(
        self,
        file_path: str,
        mode: str = "ai",
        confidence: float = 0.7,
        audit: bool = False,
        filter_config: Optional[Dict[str, bool]] = None,
    ) -> Dict[str, Any]:
        """Traitement principal avec option de rapport d'audit"""
        import time
        start_time = time.time()

        original_config = self.filter_config.copy()
        if filter_config:
            self.filter_config.update(filter_config)
            if self.ai_anonymizer:
                self.ai_anonymizer.filter_config = self.filter_config

        try:
            # Validation des paramètres
            if mode not in ["regex", "ai"]:
                mode = "regex"
            
            if not (0.0 <= confidence <= 1.0):
                confidence = 0.7
            
            # Forcer regex si IA non disponible
            if mode == "ai" and not self.ai_anonymizer:
                mode = "regex"
                logging.warning("Mode IA demandé mais non disponible, fallback vers regex")
            
            # Extraction du texte
            logging.info(f"Traitement du document: {file_path}")
            text, metadata = self.document_processor.process_file(file_path)
            
            if not text.strip():
                return {
                    "status": "error",
                    "error": "Aucun texte trouvé dans le document"
                }
            
            # Prétraitement du texte
            text = self._preprocess_text(text)
            logging.info(f"Texte extrait: {len(text)} caractères")
            
            # Détection des entités selon le mode
            if mode == "ai" and self.ai_anonymizer:
                logging.info("Détection IA en cours...")
                entities = self.ai_anonymizer.detect_entities_ai(text, confidence)
                metadata["detection_method"] = "ai"
            else:
                logging.info("Détection regex en cours...")
                entities = self.regex_anonymizer.detect_entities(text)
                metadata["detection_method"] = "regex"
            
            logging.info(f"Entités détectées: {len(entities)}")
            
            # Post-traitement et validation
            entities = self._post_process_entities(entities, text)
            
            # Anonymisation
            anonymized_text, entity_mapping = self.regex_anonymizer.anonymize_text(
                text, entities
            )
            # Conserver le mapping pour utilisation ultérieure
            self.entity_mapping = entity_mapping
            # Rendre les compteurs accessibles au niveau du document
            self.entity_counters = self.regex_anonymizer.entity_counters

            # Validation de l'anonymisation
            validation_result = self._validate_anonymization(
                text, anonymized_text, entities
            )

            # Calcul des métriques
            processing_time = time.time() - start_time
            stats = self._generate_processing_stats(entities, text, processing_time)
            metadata.update(stats)

            # Mise à jour des statistiques globales
            self.processing_stats["documents_processed"] += 1
            self.processing_stats["entities_detected"] += len(entities)
            self.processing_stats["processing_time"] += processing_time

            # Création du document anonymisé
            anonymized_path = self._create_anonymized_document(
                file_path,
                anonymized_text,
                metadata,
                entities,
                audit=audit,
            )

            return {
                "status": "success",
                "entities": [asdict(entity) for entity in entities],
                "text": text,
                "anonymized_text": anonymized_text,
                "anonymized_path": anonymized_path,
                "entity_mapping": entity_mapping,
                "entity_counters": self.entity_counters,
                "metadata": metadata,
                "mode": mode,
                "confidence": confidence,
                "validation": validation_result,
                "processing_time": processing_time,
            }
        
        except (OSError, ValueError, RuntimeError) as e:
            processing_time = time.time() - start_time
            logging.error(f"Erreur traitement document: {str(e)}")
            return {
                "status": "error",
                "error": str(e),
                "processing_time": processing_time
            }
        finally:
            # Restaurer la configuration initiale
            self.filter_config = original_config
            if self.ai_anonymizer:
                self.ai_anonymizer.filter_config = self.filter_config

    def _generate_processing_stats(
        self,
        entities: List[Entity],
        text: str,
        processing_time: float,
        thresholds: Optional[Dict[str, float]] = None,
    ) -> Dict[str, Any]:
        """Générer des statistiques complètes de traitement"""
        entity_types: Dict[str, int] = {}
        confidence_values: List[float] = []
        method_counts: Dict[str, int] = {}

        for entity in entities:
            # Types
            entity_type = entity.type
            entity_types[entity_type] = entity_types.get(entity_type, 0) + 1

            # Confiance
            if hasattr(entity, 'confidence') and entity.confidence is not None:
                confidence_values.append(entity.confidence)

            # Méthodes
            method = getattr(entity, 'method', 'unknown')
            method_counts[method] = method_counts.get(method, 0) + 1

        # Statistiques de confiance
        confidence_stats: Dict[str, Any] = {}
        thresholds = thresholds or {"high": 0.8, "medium": 0.5}
        if confidence_values:
            high = thresholds.get("high", 0.8)
            medium = thresholds.get("medium", 0.5)
            confidence_stats = {
                "min": min(confidence_values),
                "max": max(confidence_values),
                "average": sum(confidence_values) / len(confidence_values),
                "std": self._calculate_std(confidence_values),
                "high_confidence_count": len([c for c in confidence_values if c >= high]),
                "medium_confidence_count": len([c for c in confidence_values if medium <= c < high]),
                "low_confidence_count": len([c for c in confidence_values if c < medium])
            }

        return {
            "total_entities": len(entities),
            "entity_types": entity_types,
            "method_distribution": method_counts,
            "confidence_stats": confidence_stats,
            "processing_time": processing_time,
            "entities_per_second": len(entities) / processing_time if processing_time > 0 else 0,
            "text_length": len(text),
            "most_common_type": max(entity_types, key=entity_types.get) if entity_types else None,
            "ai_used": any(method in ['spacy', 'transformers'] for method in method_counts.keys())
        }

    def _calculate_std(self, values: List[float]) -> float:
        """Calculer l'écart-type"""
        if len(values) < 2:
            return 0.0

        mean = sum(values) / len(values)
        variance = sum((x - mean) ** 2 for x in values) / (len(values) - 1)
        return variance ** 0.5

    def export_anonymized_document(
        self,
        original_path: str,
        entities: Optional[List[Dict]] = None,
        options: Optional[Dict[str, Any]] = None,
        audit: bool = False,
    ) -> str:
        """Exporter un document anonymisé dans le format souhaité."""

        if (
            not isinstance(original_path, str)
            or not original_path
            or not os.path.exists(original_path)
        ):
            raise ValueError("original_path is required for export")

        options = options or {}
        output_format = options.get("format", "txt").lower()

        text, metadata = self.document_processor.process_file(original_path)

        entity_objects: List[Entity] = []
        if entities:
            for ent in entities:
                if isinstance(ent, Entity):
                    entity_objects.append(ent)
                elif isinstance(ent, dict):
                    entity_objects.append(
                        Entity(
                            id=ent.get("id", str(uuid.uuid4())),
                            type=ent.get("type", ""),
                            value=ent.get("value", ""),
                            start=ent.get("start", 0),
                            end=ent.get("end", 0),
                            confidence=ent.get("confidence", 1.0),
                            replacement=ent.get("replacement"),
                            page=ent.get("page"),
                            context=ent.get("context"),
                            method=ent.get("method", "regex"),
                            source_model=ent.get("source_model"),
                        )
                    )
                else:
                    raise ValueError("Invalid entity format")
            entities = [asdict(e) for e in entity_objects]
        else:
            detected = self.regex_anonymizer.detect_entities(text)
            entity_objects = detected
            entities = [asdict(e) for e in detected]

        anonymized_text, mapping = self.regex_anonymizer.anonymize_text(
            text, entity_objects
        )
        # Conserver le mapping pour export ultérieur
        self.entity_mapping = mapping

        if audit:
            serialized = serialize_entity_mapping(self.entity_mapping)
            if serialized is not None:
                metadata["entity_mapping"] = serialized

        stats = generate_anonymization_stats(entities, len(text)) if audit else None

        return self._write_export(
            anonymized_text,
            output_format,
            options,
            stats,
            entities,
            metadata,
            audit=audit,
            original_path=original_path,
        )

    def _create_anonymized_document(
        self,
        original_path: str,
        anonymized_text: str,
        metadata: Dict,
        entities: Optional[List[Entity]] = None,
        export_format: str = "txt",
        watermark: Optional[str] = None,
        audit: bool = False,
    ) -> str:
        """Créer un document anonymisé dans différents formats.

        Parameters
        ----------
        original_path: str
            Chemin vers le document original.
        anonymized_text: str
            Texte anonymisé à exporter.
        metadata: Dict
            Métadonnées d'anonymisation.
        entities: Optional[List[Entity]]
            Liste d'entités détectées.
        export_format: str
            Format d'export (txt, docx, pdf - nécessite le paquet fpdf).
        watermark: Optional[str]
            Filigrane à ajouter au document.
        audit: bool
            Inclure un rapport d'audit avec métadonnées et statistiques.
        """

        export_format = (export_format or "txt").lower()
        entities = entities or []
        metadata = metadata or {}

        stats: Optional[Dict[str, Any]] = None
        if audit:
            try:
                stats = generate_anonymization_stats(
                    [asdict(e) if isinstance(e, Entity) else e for e in entities],
                    len(anonymized_text),
                )
            except Exception:
                stats = None

            # Ajouter le mapping sérialisé pour le rapport d'audit
            serialized = serialize_entity_mapping(self.entity_mapping)
            if serialized is not None:
                metadata = dict(metadata)
                metadata["entity_mapping"] = serialized

        if export_format == "docx":
            if not DOCX_SUPPORT:
                raise RuntimeError("DOCX export requires python-docx")

            try:
                doc = Document(original_path)

                # Construire la table de remplacement à partir du mapping existant
                replacement_map: Dict[str, str] = {}
                if self.entity_mapping:
                    for type_map in self.entity_mapping.values():
                        for info in type_map.values():
                            token = info.get("token")
                            for variant in info.get("variants", []):
                                replacement_map[variant] = token
                elif entities:
                    for ent in entities:
                        if getattr(ent, "replacement", None):
                            replacement_map[ent.value] = ent.replacement

                # Utilitaire de remplacement simple
                def _replace_text(text: str) -> str:
                    if not text:
                        return text
                    for original, token in replacement_map.items():
                        if original in text:
                            text = text.replace(original, token)
                    return text

                def _replace_in_paragraph(paragraph):
                    for run in paragraph.runs:
                        run.text = _replace_text(run.text)

                # Remplacement basique pour le corps du document
                for paragraph in doc.paragraphs:
                    _replace_in_paragraph(paragraph)

                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                _replace_in_paragraph(paragraph)

                # En-têtes et pieds de page
                for section in doc.sections:
                    if watermark:
                        if section.header.paragraphs:
                            section.header.paragraphs[0].text = str(watermark)
                        else:
                            section.header.add_paragraph(str(watermark))

                    for paragraph in section.header.paragraphs:
                        _replace_in_paragraph(paragraph)
                    for table in section.header.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for paragraph in cell.paragraphs:
                                    _replace_in_paragraph(paragraph)

                    for paragraph in section.footer.paragraphs:
                        _replace_in_paragraph(paragraph)
                    for table in section.footer.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for paragraph in cell.paragraphs:
                                    _replace_in_paragraph(paragraph)

                # Parcours supplémentaire des parties XML (zones de texte, formes, notes, commentaires)
                WORD_NS = {
                    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                }

                def _replace_in_element(element):
                    for t in element.xpath(".//w:t", namespaces=WORD_NS):
                        t.text = _replace_text(t.text)

                # Corps principal (inclut les zones de texte et formes)
                _replace_in_element(doc.element.body)

                # Autres parties de l'histoire du document
                for part in getattr(doc.part, "story_parts", []):
                    _replace_in_element(part.element)

                # Formes inline
                for shape in getattr(doc, "inline_shapes", []):
                    _replace_in_element(shape._inline)

                # Notes de bas de page / commentaires si non couverts
                if hasattr(doc.part, "footnotes_part"):
                    _replace_in_element(doc.part.footnotes_part.element)
                if hasattr(doc.part, "endnotes_part"):
                    _replace_in_element(doc.part.endnotes_part.element)
                if hasattr(doc.part, "comments_part"):
                    _replace_in_element(doc.part.comments_part.element)

                # Métadonnées et propriétés personnalisées
                try:
                    core = doc.core_properties
                    for attr in (
                        "title",
                        "subject",
                        "creator",
                        "last_modified_by",
                        "keywords",
                        "category",
                        "comments",
                    ):
                        val = getattr(core, attr)
                        if isinstance(val, str):
                            setattr(core, attr, _replace_text(val))
                except Exception:
                    pass

                try:
                    custom_props = doc.custom_properties  # type: ignore[attr-defined]
                    for name in list(custom_props):
                        val = custom_props[name]
                        if isinstance(val, str):
                            custom_props[name] = _replace_text(val)
                except Exception:
                    pass

                # Ajout optionnel de métadonnées et statistiques dans le document
                if audit and metadata:
                    doc.add_page_break()
                    doc.add_paragraph("=== AUDIT REPORT ===")
                    for key, value in metadata.items():
                        doc.add_paragraph(f"{key}: {value}")
                    if stats:
                        doc.add_page_break()
                        doc.add_paragraph("=== STATISTICS ===")
                        for key, value in stats.items():
                            doc.add_paragraph(f"{key}: {value}")

                output_path = os.path.join(
                    self.temp_dir, f"anonymized_{uuid.uuid4().hex[:8]}.docx"
                )
                doc.save(output_path)
                return output_path

            except (OSError, ValueError) as e:
                logging.error(f"Erreur création document: {e}")
                raise RuntimeError(
                    f"Impossible de créer le document anonymisé: {str(e)}"
                ) from e

        elif export_format == "txt":
            output_path = os.path.join(
                self.temp_dir, f"anonymized_{uuid.uuid4().hex[:8]}.txt"
            )
            with open(output_path, "w", encoding="utf-8") as f:
                if watermark:
                    f.write(f"{watermark}\n\n")
                f.write(anonymized_text)
                if audit and metadata:
                    f.write("\n\n=== AUDIT REPORT ===\n")
                    for key, value in metadata.items():
                        f.write(f"{key}: {value}\n")
                    if stats:
                        f.write("\n\n=== STATISTICS ===\n")
                        for key, value in stats.items():
                            f.write(f"{key}: {value}\n")
            return output_path

        elif export_format == "pdf":
            try:
                from fpdf import FPDF
            except ImportError as e:
                raise RuntimeError(
                    "PDF export requires the 'fpdf' package. Install it with 'pip install fpdf'."
                ) from e

            output_path = os.path.join(
                self.temp_dir, f"anonymized_{uuid.uuid4().hex[:8]}.pdf"
            )
            pdf = FPDF()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            if watermark:
                pdf.set_text_color(200, 200, 200)
                pdf.text(10, 10, str(watermark))
                pdf.set_text_color(0, 0, 0)
            for line in anonymized_text.splitlines():
                pdf.multi_cell(0, 10, line)
            if audit and metadata:
                pdf.add_page()
                pdf.set_font("Arial", "B", 12)
                pdf.cell(0, 10, "AUDIT REPORT", ln=True)
                pdf.set_font("Arial", size=12)
                for key, value in metadata.items():
                    pdf.multi_cell(0, 10, f"{key}: {value}")
                if stats:
                    pdf.add_page()
                    pdf.set_font("Arial", "B", 12)
                    pdf.cell(0, 10, "STATISTICS", ln=True)
                    pdf.set_font("Arial", size=12)
                    for key, value in stats.items():
                        pdf.multi_cell(0, 10, f"{key}: {value}")
            pdf.output(output_path)
            return output_path

        else:
            raise ValueError(f"Unsupported export format: {export_format}")

    def _write_export(
        self,
        anonymized_text: str,
        output_format: str,
        options: Dict[str, Any],
        stats: Optional[Dict[str, Any]] = None,
        entities: Optional[List[Dict]] = None,
        metadata: Optional[Dict[str, Any]] = None,
        audit: bool = False,
        original_path: Optional[str] = None,
    ) -> str:
        """Écrire le texte anonymisé dans le format choisi."""

        filename = f"anonymized_{uuid.uuid4().hex[:8]}.{output_format}"
        output_path = os.path.join(self.temp_dir, filename)

        watermark = options.get("watermark")

        if output_format == "txt":
            with open(output_path, "w", encoding="utf-8") as f:
                if watermark:
                    f.write(f"{watermark}\n\n")
                f.write(anonymized_text)
                if audit:
                    if metadata:
                        f.write("\n\n=== AUDIT REPORT ===\n")
                        for key, value in metadata.items():
                            f.write(f"{key}: {value}\n")
                    if stats:
                        f.write("\n\n=== STATISTICS ===\n")
                        for key, value in stats.items():
                            f.write(f"{key}: {value}\n")
            return output_path

        elif output_format == "docx":
            if not DOCX_SUPPORT:
                raise Exception("Export DOCX non supporté")
            if not original_path:
                raise ValueError("original_path is required for DOCX export")

            doc = Document(original_path)

            # Build replacement mapping from entity_mapping or entities
            replacement_map: Dict[str, str] = {}
            if self.entity_mapping:
                for mapping in self.entity_mapping.values():
                    for info in mapping.values():
                        token = info.get("token")
                        for variant in info.get("variants", []):
                            replacement_map[variant] = token
            elif entities:
                for ent in entities:
                    if isinstance(ent, dict):
                        value = ent.get("value")
                        replacement = ent.get("replacement")
                    else:
                        value = getattr(ent, "value", None)
                        replacement = getattr(ent, "replacement", None)
                    if value and replacement:
                        replacement_map[value] = replacement

            def _replace_in_runs(runs):
                for run in runs:
                    for original, token in replacement_map.items():
                        if original in run.text:
                            run.text = run.text.replace(original, token)

            # Replace in body paragraphs
            for paragraph in doc.paragraphs:
                _replace_in_runs(paragraph.runs)

            # Replace in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            _replace_in_runs(paragraph.runs)

            # Headers and footers
            for section in doc.sections:
                if watermark:
                    if section.header.paragraphs:
                        section.header.paragraphs[0].text = str(watermark)
                    else:
                        section.header.add_paragraph(str(watermark))

                for paragraph in section.header.paragraphs:
                    _replace_in_runs(paragraph.runs)
                for table in section.header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                _replace_in_runs(paragraph.runs)

                for paragraph in section.footer.paragraphs:
                    _replace_in_runs(paragraph.runs)
                for table in section.footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                _replace_in_runs(paragraph.runs)

            # Text boxes and shapes
            try:
                from docx.oxml.ns import nsmap

                for text_elem in doc.part.element.xpath(
                    './/w:txbxContent//w:t | .//w:drawing//w:t', namespaces=nsmap
                ):
                    for original, token in replacement_map.items():
                        if original in text_elem.text:
                            text_elem.text = text_elem.text.replace(original, token)
            except Exception:
                pass

            # Append audit information if requested
            if audit:
                if metadata:
                    doc.add_paragraph("=== AUDIT REPORT ===")
                    for key, value in metadata.items():
                        doc.add_paragraph(f"{key}: {value}")
                if stats:
                    doc.add_paragraph("=== STATISTICS ===")
                    for key, value in stats.items():
                        doc.add_paragraph(f"{key}: {value}")

            output_path = options.get("output_path", original_path)
            doc.save(output_path)
            return output_path

        elif output_format == "pdf":
            try:
                from fpdf import FPDF
            except ImportError as e:
                raise RuntimeError(
                    "PDF export requires the 'fpdf' package. Install it with 'pip install fpdf'."
                ) from e

            pdf = FPDF()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            if watermark:
                pdf.set_text_color(200, 200, 200)
                pdf.text(10, 10, str(watermark))
                pdf.set_text_color(0, 0, 0)
            for line in anonymized_text.splitlines():
                pdf.multi_cell(0, 10, line)
            if audit:
                if metadata:
                    pdf.add_page()
                    pdf.set_font("Arial", 'B', 12)
                    pdf.cell(0, 10, "AUDIT REPORT", ln=True)
                    pdf.set_font("Arial", size=12)
                    for key, value in metadata.items():
                        pdf.multi_cell(0, 10, f"{key}: {value}")
                if stats:
                    pdf.add_page()
                    pdf.set_font("Arial", 'B', 12)
                    pdf.cell(0, 10, "STATISTICS", ln=True)
                    pdf.set_font("Arial", size=12)
                    for key, value in stats.items():
                        pdf.multi_cell(0, 10, f"{key}: {value}")
            pdf.output(output_path)
            return output_path

        else:
            raise ValueError(f"Unsupported format: {output_format}")

    def _preprocess_text(self, text: str) -> str:
        """Prétraitement optimisé du texte français"""
        # Normalisation des espaces
        text = re.sub(r'\s+', ' ', text)
        
        # Normalisation caractères français
        replacements = {
            'œ': 'oe', 'Œ': 'OE',
            'æ': 'ae', 'Æ': 'AE',
            '«': '"', '»': '"',
            '“': '"', '”': '"',
            '‘': "'", '’': "'",
        }
        
        for old, new in replacements.items():
            text = text.replace(old, new)
        
        # Nettoyage caractères de contrôle
        text = ''.join(char for char in text if ord(char) >= 32 or char in '\n\t')
        
        return text.strip()
    
    def _post_process_entities(self, entities: List[Entity], text: str) -> List[Entity]:
        """Post-traitement avec optimisations françaises"""
        processed = []

        for entity in entities:
            # Validation de base
            if not self._is_valid_entity(entity, text):
                continue

            # Nettoyage de la valeur
            entity.value = self._clean_entity_value(entity.value)

            # Amélioration du contexte
            if not entity.context:
                entity.context = self._extract_context(text, entity.start, entity.end)

            # Classification fine française
            entity.type = self._refine_french_classification(entity, text)

            # Filtres configurables
            cfg = self.filter_config

            if cfg.get("min_length", True):
                min_len = MIN_ENTITY_LENGTH.get(entity.type, 1)
                if len(entity.value) < min_len:
                    continue

            if cfg.get("stopwords", True) and entity.type == "PERSON":
                if entity.value.lower() in FRENCH_STOP_WORDS:
                    continue

            if cfg.get("capitalization", True) and entity.type == "PERSON":
                if not entity.value[:1].isupper():
                    continue

            if cfg.get("title_check", True) and entity.type == "PERSON":
                preceding = get_preceding_token(text, entity.start)
                if preceding.lower() in FRENCH_TITLES:
                    pass
                elif cfg.get("require_title", False):
                    continue

            processed.append(entity)

        # Résolution des conflits
        processed = self._resolve_entity_conflicts(processed)

        return processed
    
    def _is_valid_entity(self, entity: Entity, text: str) -> bool:
        """Validation robuste des entités"""
        # Vérifications de base
        if entity.start < 0 or entity.end > len(text):
            return False
        
        if entity.start >= entity.end:
            return False
        
        if len(entity.value.strip()) < 2:
            return False
        
        # Vérification cohérence texte
        actual_value = text[entity.start:entity.end]
        if actual_value.strip().lower() != entity.value.strip().lower():
            # Tentative de correction
            search_start = max(0, entity.start - 10)
            search_end = min(len(text), entity.end + 10)
            search_area = text[search_start:search_end]
            
            if entity.value in search_area:
                corrected_start = search_area.find(entity.value) + search_start
                entity.start = corrected_start
                entity.end = corrected_start + len(entity.value)
            else:
                return False
        
        return True
    
    def _clean_entity_value(self, value: str) -> str:
        """Nettoyage avancé des valeurs d'entités"""
        # Supprimer espaces début/fin
        cleaned = value.strip()
        
        # Supprimer ponctuation finale
        cleaned = re.sub(r'[.,;:!?]+$', '', cleaned)
        
        # Supprimer caractères bizarres
        cleaned = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', cleaned)
        
        # Normaliser espaces internes
        cleaned = re.sub(r'\s+', ' ', cleaned)
        
        return cleaned
    
    def _extract_context(self, text: str, start: int, end: int, context_length: int = 120) -> str:
        """Extraction de contexte intelligent"""
        context_start = max(0, start - context_length)
        context_end = min(len(text), end + context_length)
        
        # Essayer de couper à des limites de mots
        context = text[context_start:context_end]
        
        # Marquer l'entité
        entity_value = text[start:end]
        relative_start = start - context_start
        relative_end = end - context_start
        
        if 0 <= relative_start < len(context) and 0 <= relative_end <= len(context):
            highlighted_context = (
                context[:relative_start] + 
                f"**{entity_value}**" + 
                context[relative_end:]
            )
        else:
            highlighted_context = context
        
        return highlighted_context.strip()
    
    def _refine_french_classification(self, entity: Entity, text: str) -> str:
        """Classification fine pour le contexte français"""
        value = entity.value.lower()
        
        # Amélioration PERSON
        if entity.type == 'PERSON':
            context_before = text[max(0, entity.start - 30):entity.start].lower()
            
            # Vérifier titres français
            french_titles = ['monsieur', 'madame', 'mademoiselle', 'docteur', 'professeur', 'maître']
            if any(title in context_before for title in french_titles):
                return 'PERSON'
            
            # Détecter organisations
            org_words = ['société', 'entreprise', 'sarl', 'sas', 'cabinet', 'étude', 'bureau']
            if any(word in value for word in org_words):
                return 'ORG'
        
        # Amélioration ORG
        elif entity.type == 'ORG':
            legal_forms = ['sarl', 'sas', 'sa', 'snc', 'eurl', 'sasu', 'sci', 'selarl', 'selas', 'selca']
            if any(form in value for form in legal_forms):
                return 'ORG'
        
        return entity.type
    
    def _resolve_entity_conflicts(self, entities: List[Entity]) -> List[Entity]:
        """Résolution intelligente des conflits"""
        if not entities:
            return entities
        
        # Trier par position
        sorted_entities = sorted(entities, key=lambda x: x.start)
        resolved = []
        
        for current in sorted_entities:
            conflict_found = False
            
            for i, existing in enumerate(resolved):
                if self._entities_overlap(current, existing):
                    winner = self._resolve_conflict(current, existing)
                    resolved[i] = winner
                    conflict_found = True
                    break
            
            if not conflict_found:
                resolved.append(current)
        
        return resolved

    def _entities_overlap(self, entity1: Entity, entity2: Entity) -> bool:
        return not (entity1.end <= entity2.start or entity2.end <= entity1.start)

    def _resolve_conflict(self, entity1: Entity, entity2: Entity) -> Entity:
        """Résoudre les conflits en réutilisant les règles de priorité.

        Cette implémentation délègue la résolution à l'anonymiseur regex
        afin de garantir une cohérence des priorités entre les différents
        anonymiseurs. Les règles incluent la priorité des types structurés,
        de la méthode de détection, du niveau de confiance et de la longueur
        de l'entité.
        """

        # Délégation directe si possible pour conserver la même logique
        if hasattr(self.regex_anonymizer, "_resolve_conflict"):
            return self.regex_anonymizer._resolve_conflict(entity1, entity2)

        # Fallback au cas où la méthode ne serait pas disponible
        structured_types = {"EMAIL", "PHONE", "IBAN", "SIRET", "SIREN", "SSN", "TVA"}

        if entity1.type in structured_types and entity2.type not in structured_types:
            return entity1
        if entity2.type in structured_types and entity1.type not in structured_types:
            return entity2

        method_priority = {"regex": 3, "spacy": 2, "transformers": 1}
        priority1 = method_priority.get(getattr(entity1, "method", ""), 0)
        priority2 = method_priority.get(getattr(entity2, "method", ""), 0)

        if priority1 != priority2:
            return entity1 if priority1 > priority2 else entity2

        conf1 = getattr(entity1, "confidence", 0.0)
        conf2 = getattr(entity2, "confidence", 0.0)
        if abs(conf1 - conf2) > 0.15:
            return entity1 if conf1 > conf2 else entity2

        len1 = entity1.end - entity1.start
        len2 = entity2.end - entity2.start
        return entity1 if len1 >= len2 else entity2


def evaluate(dataset_path: str):
    """Évalue l'anonymiseur sur un corpus annoté.

    Parameters
    ----------
    dataset_path : str
        Chemin vers un dossier contenant des fichiers JSON avec
        les clés ``tokens`` et ``labels``.

    Returns
    -------
    pandas.DataFrame
        Tableau des métriques précision, rappel et F1 pour chaque entité.
    """

    from pathlib import Path
    import json
    import pandas as pd
    from sklearn.metrics import precision_recall_fscore_support

    dataset_dir = Path(dataset_path)
    if not dataset_dir.exists():
        raise FileNotFoundError(f"Dataset not found: {dataset_path}")

    anonymizer = RegexAnonymizer(use_french_patterns=True)

    all_true: List[str] = []
    all_pred: List[str] = []

    for file_path in dataset_dir.glob("*.json"):
        with open(file_path, "r", encoding="utf-8") as f:
            doc = json.load(f)

        tokens = doc.get("tokens", [])
        true_labels = doc.get("labels", [])
        text = " ".join(tokens)

        entities = anonymizer.detect_entities(text)
        pred_labels = ["O"] * len(tokens)

        # Calcul des offsets caractères pour chaque token
        offsets = []
        pos = 0
        for tok in tokens:
            start = pos
            end = pos + len(tok)
            offsets.append((start, end))
            pos = end + 1  # espace

        for ent in entities:
            for idx, (start, end) in enumerate(offsets):
                if not (ent.end <= start or ent.start >= end):
                    pred_labels[idx] = ent.type

        all_true.extend(true_labels)
        all_pred.extend(pred_labels)

    labels = sorted({lbl for lbl in all_true if lbl != "O"})
    precision, recall, f1, support = precision_recall_fscore_support(
        all_true, all_pred, labels=labels, average=None, zero_division=0
    )

    return pd.DataFrame(
        {
            "entity": labels,
            "precision": precision,
            "recall": recall,
            "f1": f1,
            "support": support,
        }
    )
    